#!/usr/bin/env python3
"""
build_docx.py — Word document builder for web-site-to-document skill.
Uses python-docx. Adds cover page, TOC field, and formatted content.
"""

import io
import re
from datetime import datetime
from typing import Dict, List

from common import image_bytes, is_rtl
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# A4 text height is ~24.7 cm. add_picture sets the width only, so a tall
# screenshot scaled to 5.5 in wide runs metres down the page and Word clips it.
MAX_IMAGE_HEIGHT = Cm(22)


# ─── Public API ───────────────────────────────────────────────────────────────

def build_docx(pages: List[Dict], output_path: str, source_url: str, static_toc: bool = False,
               blocked_domains=None, robots_excluded=None):
    """Generate a Word document (.docx) from scraped pages.

    static_toc renders a plain contents list instead of the Word TOC field,
    for outputs converted without Word (headless LibreOffice PDF).
    blocked_domains and robots_excluded are recorded in the archive: a reader
    months later cannot otherwise tell a gap in coverage from a site that had
    nothing there.
    """
    doc = Document()
    _setup_styles(doc)
    _set_update_fields(doc)

    # A4 to match the PDF output; python-docx defaults to US Letter
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Cover page
    _add_cover(doc, source_url, len(pages))
    doc.add_page_break()

    # TOC placeholder (Word updates it automatically when opened)
    if static_toc:
        _add_static_toc(doc, pages)
    else:
        _add_toc_field(doc)
    doc.add_page_break()

    # Content
    all_binary_refs: List[Dict] = []

    for i, page in enumerate(pages):
        title = page.get("title") or f"Page {i + 1}"
        url = page.get("url", "")
        breadcrumb = page.get("breadcrumb", [])

        # Page title as Heading 1 — site text, so it takes the site's direction
        page_h = doc.add_heading(title, level=1)
        if is_rtl(title):
            _rtl_para(page_h)

        # Metadata (URL + breadcrumb)
        if url:
            p = doc.add_paragraph()
            run = p.add_run(f"Source: {url}")
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(9)

        if breadcrumb:
            p = doc.add_paragraph()
            crumb = " › ".join(breadcrumb)
            run = p.add_run(crumb)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)
            if is_rtl(crumb):
                _rtl_para(p)

        # Content blocks
        for block in page.get("content", []):
            _add_block(doc, block)

        all_binary_refs.extend(page.get("binary_references", []))

        if i < len(pages) - 1:
            doc.add_page_break()

    # References section
    if all_binary_refs or blocked_domains or robots_excluded:
        doc.add_page_break()
        doc.add_heading("References and Attachments", level=1)
        seen = set()
        for ref in all_binary_refs:
            if ref["url"] in seen:
                continue
            seen.add(ref["url"])
            p = doc.add_paragraph(style="List Bullet")
            name = ref.get("name") or ref["url"]
            ftype = ref.get("type", "FILE")
            p.add_run(f"{name} ({ftype}): ")
            run = p.add_run(ref["url"])
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
            run.font.size = Pt(9)

        if blocked_domains:
            doc.add_heading("Domains not reachable during extraction", level=2)
            for domain in sorted(blocked_domains):
                doc.add_paragraph(domain, style="List Bullet")

        if robots_excluded:
            doc.add_heading("Pages excluded by robots.txt", level=2)
            for page_url in robots_excluded:
                doc.add_paragraph(page_url, style="List Bullet")

    doc.save(output_path)


# ─── Styles ───────────────────────────────────────────────────────────────────

def _setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Code Block style. The default template has no style by that name, so the
    # add always succeeds; _add_block still guards its use of the style.
    cs = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    cs.font.name = "Courier New"
    cs.font.size = Pt(9)
    cs.paragraph_format.left_indent = Cm(1)
    cs.paragraph_format.space_before = Pt(4)
    cs.paragraph_format.space_after = Pt(4)


# ─── Cover ────────────────────────────────────────────────────────────────────

def _add_cover(doc: Document, url: str, page_count: int):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Website Archive")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(url)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    doc.add_paragraph()

    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Extracted: {date_str}\nPages: {page_count}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ─── TOC ─────────────────────────────────────────────────────────────────────

def _set_update_fields(doc: Document):
    """Tell Word to update fields (fills the TOC) when the document opens."""
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _add_static_toc(doc: Document, pages: List[Dict]):
    """Plain contents list — headless conversion cannot update a TOC field."""
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.add_run("Table of Contents")
    for i, page in enumerate(pages):
        title = page.get("title") or f"Page {i + 1}"
        doc.add_paragraph(title, style="List Bullet")


def _add_toc_field(doc: Document):
    """Insert a Word TOC field that updates when the document is opened."""
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.add_run("Table of Contents")

    p = doc.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr)

    fld_end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "[Right-click → Update Field to refresh the Table of Contents in Word]"
    )
    note_run.font.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ─── Block Rendering ──────────────────────────────────────────────────────────

def _block_text(block: Dict) -> str:
    """The text whose direction decides the block's direction."""
    if block.get("items"):
        return " ".join(str(i) for i in block["items"])
    if block.get("data"):
        return " ".join(str(c) for row in block["data"] for c in row)
    return block.get("text") or block.get("alt") or ""


# w:bidi's place in the w:pPr sequence: everything below must follow it. The children of a
# pPr are an ordered sequence in the schema, not a set, and Word repairs a file whose order is
# wrong — so the element is inserted before its successors rather than appended.
_BIDI_SUCCESSORS = (
    "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing",
    "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
    "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
    "w:pPrChange",
)


def _rtl_para(p):
    """Mark one paragraph right-to-left.

    `w:bidi` is what decides which edge the paragraph starts at and which side a list number
    sits on; the run-level flag is what puts the trailing full stop on the correct side. Word
    shapes the glyphs either way, so without these the text is readable but laid out backwards.
    python-docx has no paragraph-level API for it, hence the raw element — but it does own the
    run-level one, so that half goes through `font.rtl`.
    """
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.insert_element_before(bidi, *_BIDI_SUCCESSORS)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.rtl = True
    return p


def _rtl_table(tbl):
    """Mirror a right-to-left table's columns: the first cell belongs on the right.

    Through python-docx's own accessor, which places w:bidiVisual at its schema position;
    appending it lands after w:tblW, which is out of order.
    """
    tbl._tbl.tblPr.get_or_add_bidiVisual()
    return tbl


def _add_block(doc: Document, block: Dict):
    btype = block.get("type", "")
    # Code keeps its left-to-right layout on every page: a shell command reads the same way
    # in Riyadh as in Rome.
    rtl = btype != "code" and is_rtl(_block_text(block))

    if btype == "heading":
        level = min(block.get("level", 1) + 1, 9)
        h = doc.add_heading(block.get("text", ""), level=level)
        if rtl:
            _rtl_para(h)

    elif btype == "paragraph":
        text = block.get("text", "").strip()
        if text:
            p = doc.add_paragraph()
            _inline(p, text)
            if rtl:
                _rtl_para(p)

    elif btype == "code":
        code = block.get("text", "")
        lang = block.get("language", "")
        try:
            p = doc.add_paragraph(style="Code Block")
        except Exception:
            p = doc.add_paragraph()
        if lang:
            r = p.add_run(f"[{lang}]\n")
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.font.size = Pt(8)
        r = p.add_run(code)
        r.font.name = "Courier New"
        r.font.size = Pt(9)

    elif btype == "list":
        base = "List Number" if block.get("ordered") else "List Bullet"
        # Word ships "List Bullet 2/3" and "List Number 2/3" for nesting;
        # anything deeper renders at level 3
        level = min(block.get("level", 0), 2)
        style = f"{base} {level + 1}" if level else base
        for item in block.get("items", []):
            p = doc.add_paragraph(style=style)
            _inline(p, item)
            if rtl:
                _rtl_para(p)

    elif btype == "table":
        data = block.get("data", [])
        if data:
            rows = len(data)
            cols = max(len(r) for r in data)
            if rows and cols:
                tbl = doc.add_table(rows=rows, cols=cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(data):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = tbl.cell(ri, ci)
                        _inline(cell.paragraphs[0], str(cell_text))
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                        if rtl:
                            _rtl_para(cell.paragraphs[0])
                if rtl:
                    _rtl_table(tbl)
                doc.add_paragraph()

    elif btype == "image":
        img_bytes = image_bytes(block)
        if img_bytes:
            try:
                stream = io.BytesIO(img_bytes)
                pic = doc.add_picture(stream, width=Inches(5.5))
                if pic.height > MAX_IMAGE_HEIGHT:
                    pic.width = int(pic.width * MAX_IMAGE_HEIGHT / pic.height)
                    pic.height = MAX_IMAGE_HEIGHT
                alt = block.get("alt", "")
                if alt:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(alt)
                    cr.font.italic = True
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                    if rtl:
                        _rtl_para(cp)
            except Exception:
                _image_placeholder(doc, block)
        else:
            _image_placeholder(doc, block)

    elif btype == "callout":
        text = block.get("text", "").strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            r = p.add_run(text)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if rtl:
                _rtl_para(p)


def _image_placeholder(doc: Document, block: Dict):
    # Alt text *and* URL: Word cannot embed SVG, and a diagram reduced to its
    # caption with no address left is a dead end for the reader
    alt = block.get("alt", "")
    src = block.get("src", "")
    p = doc.add_paragraph()
    r = p.add_run(f"[Image: {' — '.join(x for x in (alt, src) if x)}]")
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# The link form comes first: without it the alternation below would match the
# *italic* inside a URL and Word would print raw [label](url) brackets
LINK_RE = r"\[[^\]\[]+\]\([^()\s]+\)"


def _delink(text: str) -> str:
    """Every [label](url) in text reduced to its label.

    Word gets the label alone: python-docx has no hyperlink API, and the page's
    own address already prints under the section title. The emphasis branches
    below need this because the alternation matches "**[y](url)**" as one bold
    part, so the link branch never sees it and Word printed raw brackets.
    """
    return re.sub(LINK_RE, lambda m: m.group()[1:m.group().index("](")], text)


def _inline(p, text: str):
    """Parse simple markdown inline formatting (**bold**, *italic*, `code`).

    A [label](url) link renders as its label: python-docx has no hyperlink API,
    and the address is already on the page's Source line above.
    """
    pattern = re.compile("(" + LINK_RE + r"|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("[") and part.endswith(")"):
            # Recurse: a label may itself carry **bold**, and a label never
            # carries a second link (the extractor skips bracketed labels)
            _inline(p, part[1:part.index("](")])
        elif part.startswith("**") and part.endswith("**"):
            r = p.add_run(_delink(part[2:-2]))
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(_delink(part[1:-1]))
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            p.add_run(part)
