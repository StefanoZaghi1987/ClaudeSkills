#!/usr/bin/env python3
"""
test_scrape.py — Runnable check for the page parser and Word table rendering.
No framework, no network.

Run:  python skills/web-site-to-document/scripts/test_scrape.py
Exits non-zero on the first failed assertion.
"""

import contextlib
import io
import os
import sys
import time

# Windows pipes default to the ANSI code page, which cannot encode the glyphs
# scrape() prints when the seeding tests below run it
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import requests

from bs4 import BeautifulSoup

from crawl import MAX_PAGE_BYTES as _crawl_MAX_PAGE_BYTES
from crawl import SiteAnalyzer, WebScraper, SPARSE_PROBE_PAGES
from extract import PageExtractor

HTML = """<html><head><title>Old title</title></head><body><main>
<h1>Real title</h1>
<p>Hello <a href="/x">link</a>, <strong>bold</strong>, <em>ital</em>, and <code>code</code>.</p>
<img src="/i.png" alt="pic">
<ul><li>one <a href="/y">ref</a></li><li>two</li></ul>
<ol><li>step</li></ol>
<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td><a href="/z">detail</a></td></tr></table>
<h2>Docs <a href="/w">index</a></h2>
<pre><code class="language-python">print(1)</code></pre>
<blockquote>Note text</blockquote>
</main></body></html>"""

page = PageExtractor().extract(HTML, "https://example.com/docs/")

types = [b["type"] for b in page["content"]]
# The h1 becomes the page title (rendered as section header by the builders),
# so it must NOT reappear as a content block
assert types == ["paragraph", "image", "list", "list", "table", "heading", "code", "callout"], types
assert page["title"] == "Real title", page["title"]
assert page["content"][0]["text"] == (
    "Hello [link](https://example.com/x), **bold**, *ital*, and `code`.")
# No "data": a block records where the picture was written, never the picture
# itself. Carrying the base64 as well held every one of them in memory for the
# whole of a crawl whose page count is unlimited by default.
assert page["content"][1] == {
    "type": "image", "src": "https://example.com/i.png", "alt": "pic",
    "mime_type": None, "path": None,
}, page["content"][1]
assert page["content"][2] == {
    "type": "list", "ordered": False, "level": 0, "start": 0,
    "items": ["one [ref](https://example.com/y)", "two"]}, page["content"][2]
assert page["content"][3] == {"type": "list", "ordered": True, "level": 0,
                              "start": 0, "items": ["step"]}
assert page["content"][4] == {"type": "table", "data": [
    ["A", "B"], ["1", "[detail](https://example.com/z)"]]}, page["content"][4]
assert page["content"][5] == {"type": "heading", "level": 2, "text": "Docs index"}
assert page["content"][6] == {"type": "code", "text": "print(1)", "language": "python"}
assert page["content"][7] == {"type": "callout", "text": "Note text"}
# Links nested in headings and table cells are collected for traversal, not just p/li text
assert page["links"] == ["https://example.com/x", "https://example.com/y",
                         "https://example.com/z", "https://example.com/w"], page["links"]

# Discovery vs fidelity: sidebar links are followed, sidebar text is not
# extracted. Bare text (no <p> wrapper), <dl>, and <figcaption> are content.
HTML2 = """<html><head><title>Nav page</title></head><body>
<nav class="sidebar"><a href="/docs/a">Section A</a><a href="/docs/b">Section B</a></nav>
<main>
  <div>First bare paragraph.</div>
  <div><span>Second bare line.</span></div>
  <p>Wrapped paragraph.</p>
  <dl><dt>Term</dt><dd>Definition text</dd></dl>
  <figure><img src="/x.png" alt="pic"><figcaption>Caption text</figcaption></figure>
  <!-- a comment that must not leak -->
</main></body></html>"""

page2 = PageExtractor().extract(HTML2, "https://example.com/")
assert page2["links"] == ["https://example.com/docs/a", "https://example.com/docs/b"], page2["links"]
assert [b["type"] for b in page2["content"]] == [
    "paragraph", "paragraph", "paragraph", "paragraph", "paragraph",
    "image", "paragraph",
], page2["content"]
texts = [b.get("text") for b in page2["content"]]
assert texts[:5] == [
    "First bare paragraph.", "Second bare line.", "Wrapped paragraph.",
    "Term", "Definition text",
], texts
assert texts[-1] == "Caption text", texts
assert all("Section" not in (t or "") for t in texts), texts  # nav text is not content

# Title source: the <h1> the *cleaned* content area still holds. A site-wide
# <h1> in the page header is the site name, and taking it would title every page
# of the site identically; <title> is the fallback.
LOGO = """<html><head><title>Install guide - ACME Docs</title></head><body>
<header><h1>ACME Docs</h1></header>
<main><h2>Install guide</h2><p>Body text with well over eighty characters, so that
the content-area selector picks this main element and not the whole body.</p></main>
</body></html>"""
logo = PageExtractor().extract(LOGO, "https://acme.example/install")
assert logo["title"] == "Install guide - ACME Docs", logo["title"]
assert [b["type"] for b in logo["content"]] == ["heading", "paragraph"], logo["content"]
assert PageExtractor().extract(
    "<html><body><main><h1>Own heading</h1><p>Body text with well over eighty "
    "characters, so the content-area selector picks main.</p></main></body></html>",
    "https://acme.example/x",
)["title"] == "Own heading"

# Noise selectors match whole words: "toc" is a substring of "protocol" and of
# "in-stock", and eating those blocks silently deletes real page content
TOC = """<html><head><title>T</title></head><body><main>
<div class="protocol-spec"><p>Protocol body.</p></div>
<div class="in-stock"><p>Stock body.</p></div>
<div class="page-toc"><p>Contents noise</p></div>
<div id="toc"><p>More contents noise</p></div>
<p>Baseline paragraph long enough to make this a genuine content area.</p>
</main></body></html>"""
toc = [b.get("text") for b in PageExtractor().extract(TOC, "https://example.com/")["content"]]
assert toc[:2] == ["Protocol body.", "Stock body."], toc
assert not any("noise" in t for t in toc), toc
for _cls in ("toc", "page-toc", "toc-container", "page_toc", "toc_wrap", "table-of-contents"):
    _h = (f'<html><head><title>T</title></head><body><main><div class="{_cls}">'
          "<p>Contents noise</p></div><p>Baseline paragraph long enough to make "
          "this a genuine content area.</p></main></body></html>")
    _blocks = PageExtractor().extract(_h, "https://example.com/")["content"]
    assert not any("noise" in (b.get("text") or "") for b in _blocks), _cls
# ...but a noise selector names a block of furniture, never an inline fragment
# inside content. `<h2><a class="toc-backref">…</a></h2>` is what docutils writes
# for every section heading it can link back to, and [class*="toc-"] matched that
# anchor: the heading came out empty, _walk emitted nothing, and PEP 8 reached the
# archive without one of its 42 section titles (live-verified, 0 headings -> 42).
BACKREF = ('<html><head><title>T</title></head><body><main>'
           '<section><h2><a class="toc-backref" href="#intro">Introduction</a></h2>'
           '<p>Body text with well over eighty characters, so the content-area '
           'selector picks this main element.</p></section>'
           '</main></body></html>')
_br = PageExtractor().extract(BACKREF, "https://example.com/")["content"]
assert _br[0] == {"type": "heading", "level": 2, "text": "Introduction"}, _br
# a real contents block still goes: it is the container that matches, not an anchor
_toc_block = PageExtractor().extract(
    '<html><head><title>T</title></head><body><main>'
    '<div class="toc-container"><a href="#a">Intro</a></div>'
    '<p>Body text with well over eighty characters, so the content-area selector '
    'picks this main element.</p></main></body></html>', "https://example.com/")
assert all("Intro" not in (b.get("text") or "") for b in _toc_block["content"]), _toc_block


# Nested lists: a sublist is a block of its own carrying its level, placed right
# after the items it hangs from, and an ordered list keeps counting across the
# split. Block children of an <li> must not glue into the neighbouring words.
NESTED = """<html><body><main>
<ol><li>Parent<ul><li>alpha</li><li>beta</li></ul></li><li>two</li></ol>
<ul><li><p>one</p><p>and two</p></li></ul>
</main></body></html>"""

nested = PageExtractor().extract(NESTED, "https://example.com/")["content"]
assert nested == [
    {"type": "list", "ordered": True, "level": 0, "start": 0, "items": ["Parent"]},
    {"type": "list", "ordered": False, "level": 1, "start": 0, "items": ["alpha", "beta"]},
    {"type": "list", "ordered": True, "level": 0, "start": 1, "items": ["two"]},
    {"type": "list", "ordered": False, "level": 0, "start": 0, "items": ["one and two"]},
], nested

# An <img> nested inside a block child of a list item or paragraph is captured,
# not lost to get_text(): a card grid puts the picture in a <div> inside the <li>,
# and an inline <span> wrapper takes the same text-only path. Each nesting must
# yield exactly one image block — _walk reaches the same shapes under a plain
# container, so a second capture would duplicate every picture on the page.
PAD = ("<p>Padding paragraph long enough to make this main a genuine content "
       "area.</p>")
for _inner, _want in [
    ('<ul><li><article><div class="image_container"><a href="/b">'
     '<img src="/c.jpg" alt="Cover"></a></div><h3>Title</h3></article></li></ul>',
     ["list", "image", "paragraph"]),
    ('<p>text<span><img src="/a.jpg" alt="A"></span></p>',
     ["paragraph", "image", "paragraph"]),
    ('<ul><li><img src="/a.jpg" alt="A">cap</li></ul>',
     ["list", "image", "paragraph"]),
    ('<p><a href="/x"><img src="/a.jpg" alt="A"></a></p>',
     ["image", "paragraph"]),
]:
    _got = [b["type"] for b in PageExtractor().extract(
        f"<html><head><title>T</title></head><body><main>{_inner}{PAD}"
        "</main></body></html>", "https://example.com/")["content"]]
    assert _got == _want, (_inner, _got)

# A navigation word in a class or an id is a hint, not proof. Word-bounding the
# patterns fixed "protocol-spec" in round 7 but cannot generalise: "page-toc" and
# "lunch-menu" are the same string shape. Every KEEP row below reached the
# document as nothing at all — a whole content block deleted in silence.
#
# Link density is what separates them: navigation is mostly link text, prose is
# not. Length deliberately is not a second signal — a one-line content block and
# a one-line consent bar are the same length.
def _content_texts(inner):
    _bs = PageExtractor().extract(
        f"<html><head><title>T</title></head><body><main>{inner}{PAD}</main>"
        "</body></html>", "https://example.com/")["content"]
    _out = []
    for _b in _bs:
        _out += (_b.get("items", []) if _b.get("type") == "list"
                 else ([_b["text"]] if _b.get("text") else []))
    return [t for t in _out if "Padding paragraph" not in t]


_LONG = "This site stores cookies. " * 30
for _label, _html in [
    ("lunch menu", '<div class="lunch-menu"><p>Risotto 12 EUR</p></div>'),
    ("banner specs", '<div class="banner-specs"><p>Torque 400 Nm</p></div>'),
    ("announcements", '<div class="announcements-archive"><p>Release 2.1</p></div>'),
    ("navigation docs", '<div class="navigation-api-reference"><p>Route table</p></div>'),
    ("naval id", '<div id="naval-architecture"><p>Hull design</p></div>'),
    # an <article>'s own header carries its byline and date, not site chrome
    ("article header",
     "<article><header><h2>Post title</h2><p>By the docs team, 2026</p>"
     "</header><p>Body.</p></article>"),
    # ...and a page that IS the cookie policy keeps its body
    ("cookie policy", f'<div class="cookie-policy"><p>{_LONG}</p></div>'),
]:
    assert _content_texts(_html), f"content deleted as noise: {_label}"

for _label, _html in [
    ("sidebar", '<div class="sidebar"><h3>Guides</h3><ul><li><a href="/a">Getting '
                'started</a></li><li><a href="/b">Install</a></li></ul></div>'),
    ("page toc", '<div class="page-toc"><ul><li><a href="/a">Intro</a></li>'
                 '<li><a href="/b">Setup</a></li></ul></div>'),
    # a breadcrumb whose last crumb is the current page is mostly UNlinked text,
    # so density would keep it — measured leaking in live on books.toscrape.com.
    # "breadcrumb" has one meaning, and the extractor keeps it as its own field.
    ("breadcrumb", '<ul class="breadcrumb"><li><a href="/">Home</a></li>'
                   '<li class="active">All products</li></ul>'),
    # a page header with a long tagline is mostly unlinked too; its position,
    # not its density, is what makes it chrome
    ("site header", '<header><a href="/"><span>ACME</span></a>'
                    '<small>We love being scraped, and this line is long.</small>'
                    '</header>'),
    ("footer", '<footer><a href="/privacy">Privacy</a> '
               '<a href="/terms">Terms</a></footer>'),
    ("menu", '<div class="main-menu"><a href="/x">Home</a>'
             '<a href="/y">About</a></div>'),
    ("nav tag", '<nav><a href="/a">A</a><a href="/b">B</a></nav>'),
]:
    assert not _content_texts(_html), f"navigation kept as content: {_label}"

# <template> is inert by specification: no browser renders what it holds, so its
# row blueprints and {{placeholders}} reached the document as text the page never
# showed. A hidden element is only a hint by comparison — an accordion panel
# starts hidden and holds content the reader wants — so it deliberately stays.
_TPL = ('<div class="wrap"><p>Real body text, long enough to be a content area.'
        '</p><template><p>Row blueprint {{name}}</p></template></div>')
assert _content_texts(_TPL) == ["Real body text, long enough to be a content area."],     _content_texts(_TPL)
_HID = '<div hidden><p>Accordion panel body the reader still wants.</p></div>'
assert _content_texts(_HID) == ["Accordion panel body the reader still wants."],     _content_texts(_HID)

# Paragraphs: the whitespace the HTML source is merely indented with collapses to
# single spaces, while the line breaks <br> asked for survive
WS = """<html><body><main>
<p>
  Hello
  world   again
</p>
<p>first<br>second</p>
</main></body></html>"""
ws = PageExtractor().extract(WS, "https://example.com/")["content"]
assert ws[0]["text"] == "Hello world again", repr(ws[0]["text"])
assert ws[1]["text"] == "first\nsecond", repr(ws[1]["text"])

# bytes input: the page's <meta charset> governs decoding — when the server
# header declares no charset, resp.text would decode UTF-8 as latin-1 mojibake
page3 = PageExtractor().extract(
    '<html><head><meta charset="utf-8"></head><body><main><p>Prezzo £ 10</p></main></body></html>'.encode("utf-8"),
    "https://example.com/",
)
assert page3["content"][0]["text"] == "Prezzo £ 10", page3["content"][0]

# URL normalization: tracking params stripped, lang kept (distinct localized
# pages must not merge). __new__ skips __init__, which would fetch robots.txt.
dummy = WebScraper.__new__(WebScraper)
n = WebScraper._normalize(dummy, "https://example.com/a?utm_source=x&lang=it")
assert n == "https://example.com/a/?lang=it", n
n = WebScraper._normalize(dummy, "https://EXAMPLE.com/Page#sec")
assert n == "https://example.com/Page/", n
n = WebScraper._normalize(dummy, "https://example.com")
assert n == "https://example.com/", n  # bare domain and "/" share one visited key
assert WebScraper._normalize(dummy, "https://example.com/index.html") == "https://example.com/"
assert WebScraper._normalize(dummy, "https://example.com/catalogue/index.html") == "https://example.com/catalogue/"
assert WebScraper._normalize(dummy, "https://example.com/page.html") == "https://example.com/page.html"

# Subdomain scope: for a two-level country ending the registrable root is
# example.co.uk — not co.uk, which would match every unrelated .co.uk site
s = WebScraper.__new__(WebScraper)
s.scope = "subdomains"
s.base_domain = "help.example.co.uk"
assert WebScraper._in_scope(s, "https://docs.example.co.uk/a") is True
assert WebScraper._in_scope(s, "https://example.co.uk/") is True
assert WebScraper._in_scope(s, "https://unrelated.co.uk/") is False
s.base_domain = "help.example.com"
assert WebScraper._in_scope(s, "https://www.example.com/a") is True
assert WebScraper._in_scope(s, "https://other.com/") is False

# www folding (same scope): the apex and its www form are one site, for scope
# and for the visited-set key; an unrelated host stays out
w = WebScraper.__new__(WebScraper)
w.scope = "same"
w.base_domain = "www.example.com"
assert WebScraper._in_scope(w, "https://example.com/a") is True
assert WebScraper._in_scope(w, "https://www.other.com/") is False
w.base_domain = "help.example.com"
assert WebScraper._in_scope(w, "https://www.help.example.com/a") is True  # www of the same host folds
assert WebScraper._in_scope(w, "https://example.com/") is False  # different host of the same root
assert WebScraper._normalize(dummy, "https://www.example.com/a") == "https://example.com/a/"

# --scope custom with no --allowed-domains falls back to the "same" rule, www
# folding included. A stricter rule of its own archived one page of any site that
# mixes example.com and www.example.com in its own links.
c = WebScraper.__new__(WebScraper)
c.scope = "custom"
c.allowed_domains = []
c.base_domain = "www.example.com"
assert WebScraper._in_scope(c, "https://example.com/a") is True
assert WebScraper._in_scope(c, "https://other.com/a") is False

# Sitemap seeding is gated on unlimited depth: a depth-limited run follows
# links only. _fetch_page is stubbed — no network.
def _no_net_scraper(max_depth):
    x = WebScraper.__new__(WebScraper)
    x.base_url = "https://example.com/"
    x.base_domain = "example.com"
    x.scope = "same"
    x.allowed_domains = []
    x.max_depth = max_depth
    x.max_pages = None
    x.rate_limit = 0
    x._robots_by_host = {"example.com": None}   # no robots.txt, no fetch
    x.js_rendered = False
    x.visited = set()
    x.blocked_domains = set()
    x.robots_excluded = []
    x.consecutive_errors = 0
    x.total_errors = 0
    x._fetch_page = lambda url, depth: {
        "url": url, "title": url, "content": [], "links": [], "binary_references": [],
    }
    return x

pages = _no_net_scraper(2).scrape(sitemap_urls=["https://example.com/from-sitemap"])
assert [p["url"] for p in pages] == ["https://example.com/"], pages
pages = _no_net_scraper(None).scrape(sitemap_urls=["https://example.com/from-sitemap"])
assert sorted(p["url"] for p in pages) == ["https://example.com/", "https://example.com/from-sitemap"], pages

# Untrusted hrefs: a malformed one (urlparse raises on unmatched IPv6 brackets)
# is skipped, and never takes the page or the site analysis down with it
BAD = ('<html><head><title>T</title></head><body><main>'
       '<a href="http://[::1">bad</a><a href="/good">good</a>'
       '<p>Body text with enough characters to look like a real content area.</p>'
       '</main></body></html>')
bad_page = PageExtractor().extract(BAD, "https://example.com/")
assert bad_page["links"] == ["https://example.com/good"], bad_page["links"]
assert SiteAnalyzer._count_links(BeautifulSoup(BAD, "lxml"), "https://example.com/") == 1

# An unparseable sitemap <loc> is skipped, not fatal
seeded = _no_net_scraper(None).scrape(sitemap_urls=["http://[::1", "https://example.com/ok"])
assert sorted(p["url"] for p in seeded) == ["https://example.com/", "https://example.com/ok"], seeded

# Redirects: the URL the response actually came from joins the visited set, so a
# direct link to it is not archived a second time.
r = _no_net_scraper(None)
r._fetch_page = lambda url, depth: {
    "url": "https://example.com/final", "title": url, "content": [],
    "links": [], "binary_references": [],
}
pages = r.scrape(sitemap_urls=["https://example.com/a", "https://example.com/b"])
assert len(pages) == 1, [p["title"] for p in pages]


# A domain that refuses connections is excluded after one failure; a connect
# timeout is a per-request failure and must never write off the whole domain.
class _RaisingSession:
    def __init__(self, exc):
        self.exc, self.calls = exc, 0

    def get(self, url, **kw):
        self.calls += 1
        raise self.exc


def _net_scraper(exc):
    x = _no_net_scraper(None)
    del x._fetch_page  # exercise the real one, against a session that only raises
    x.session = _RaisingSession(exc)
    return x

s = _net_scraper(requests.exceptions.ConnectionError("refused"))
assert s.scrape(sitemap_urls=["https://example.com/a", "https://example.com/b"]) == []
assert s.blocked_domains == {"example.com"}, s.blocked_domains
assert s.session.calls == 1, s.session.calls   # excluded, not retried per URL
assert s.total_errors == 0, s.total_errors     # a dead domain is not a page error

s = _net_scraper(requests.exceptions.ConnectTimeout("slow"))
assert s.scrape(sitemap_urls=["https://example.com/a", "https://example.com/b"]) == []
assert s.blocked_domains == set(), s.blocked_domains
assert s.session.calls == 3, s.session.calls
assert s.total_errors == 3, s.total_errors

# Non-HTML responses are skipped, not counted as errors
class _CtSession:
    def __init__(self):
        self.read = False

    def get(self, url, **kw):
        outer = self

        class R:
            status_code = 200
            headers = {"content-type": "application/json"}

            def close(self):
                pass

            def iter_content(self, n):
                outer.read = True     # a non-HTML body must never be read
                yield b"{}"
        return R()

s = _no_net_scraper(None)
del s._fetch_page
s.session = _CtSession()
assert s.scrape() == []
assert s.total_errors == 0
# ...and the body is never read. Without stream=True requests read every body in
# full first, so a thumbnail linking to its own full-size JPEG had that JPEG
# downloaded and thrown away, once per picture, at the crawl's pace.
assert s.session.read is False, "a non-HTML body was downloaded to be discarded"


# A Content-Type charset is page-authored, so it can name a codec Python has not
# got. requests.text guards that; decoding by hand has to guard it too, or the
# page arrives as an error instead of as its bytes.
class _OddCharsetSession:
    def get(self, url, **kw):
        class R:
            status_code = 200
            headers = {"content-type": "text/html; charset=utf8x"}
            encoding = "utf8x"
            url = "https://example.com/"

            def close(self):
                pass

            def iter_content(self, n):
                yield b"<html><body><main><p>Body text here.</p></main></body></html>"
        return R()


_odd = _no_net_scraper(None)
del _odd._fetch_page
_odd.session = _OddCharsetSession()
_odd.extractor = PageExtractor()
_odd_pages = _odd.scrape()
assert [b["text"] for b in _odd_pages[0]["content"]] == ["Body text here."], _odd_pages
assert _odd.total_errors == 0, _odd.total_errors


# A page is untrusted input like an image or a compressed sitemap: read to the
# cap, not past it, or one endless response decides the run's memory use
class _EndlessSession:
    def __init__(self):
        self.sent = 0

    def get(self, url, **kw):
        outer = self

        class R:
            status_code = 200
            headers = {"content-type": "text/html"}

            def close(self):
                pass

            def iter_content(self, n):
                # Endless, but not literally: a cap that stopped working has to
                # fail this check rather than hang it, so the stub stops itself
                # at twice what the reader is allowed to take.
                while outer.sent < 2 * _crawl_MAX_PAGE_BYTES:
                    outer.sent += n
                    yield b"x" * n
        return R()


_big = _no_net_scraper(None)
del _big._fetch_page
_big.session = _EndlessSession()
assert _big.scrape() == []
assert _big.total_errors == 0, "an oversized page is a skip, not an error"
assert _big.session.sent <= _crawl_MAX_PAGE_BYTES + 64 * 1024, _big.session.sent

# ─── The fetch path against a real server ─────────────────────────────────────
# Every check above hands WebScraper a session written by hand, so the parts
# requests actually performs — streaming, iter_content, close() on a body nobody
# read, a connection reused across URLs — are exactly the parts no stub can get
# wrong. This one runs them, offline, against a loopback server.
import http.server
import threading

_LOOP_SENT = {"jpeg": 0, "endless": 0}
_LOOP_HITS = {}
_JPEG_BYTES = 32 * 1024 * 1024
_LOOP_CHUNK = 64 * 1024


class _LoopHandler(http.server.BaseHTTPRequestHandler):
    protocol_version = "HTTP/1.1"          # keep-alive, so a connection is reused

    def log_message(self, *a):
        pass                               # the crawl's own output is this check's

    def _head(self, status, ctype, length):
        self.send_response(status)
        self.send_header("content-type", ctype)
        self.send_header("content-length", str(length))
        self.end_headers()

    def _pour(self, key, total):
        """Write up to `total` bytes, counting what went out before the hang-up."""
        chunk = b"x" * _LOOP_CHUNK
        try:
            while _LOOP_SENT[key] < total:
                self.wfile.write(chunk)
                _LOOP_SENT[key] += _LOOP_CHUNK
        except OSError:
            pass                           # the client closed: that is the point

    def do_GET(self):
        route = self.path.rstrip("/") or "/"
        _LOOP_HITS[route] = _LOOP_HITS.get(route, 0) + 1
        if route == "/":
            body = (b"<html><body><main><p>" + b"Body text here. " * 20 +
                    b'</p><p><a href="/big.jpg">thumb</a>'
                    b'<a href="/endless">more</a></p></main></body></html>')
            self._head(200, "text/html; charset=utf-8", len(body))
            self.wfile.write(body)
        elif route == "/big.jpg":
            # .jpg is no BINARY_EXTENSION, so this link is followed like any
            # other: the thumbnail pointing at its own full-size picture, which
            # requests used to fetch whole before the content type was read
            self._head(200, "image/jpeg", _JPEG_BYTES)
            self._pour("jpeg", _JPEG_BYTES)
        elif route == "/endless":
            # Declares far more than the page cap, and means it
            self._head(200, "text/html", 3 * _crawl_MAX_PAGE_BYTES)
            self._pour("endless", 2 * _crawl_MAX_PAGE_BYTES)
        else:
            self._head(404, "text/plain", 0)     # /robots.txt lands here


class _LoopServer(http.server.ThreadingHTTPServer):
    def handle_error(self, request, client_address):
        # Abandoning a body mid-stream is the behaviour under test, and it
        # reaches the server as a reset on the next keep-alive read. Left to
        # print, it buries the run in two 20-line tracebacks. Only the socket
        # errors go quiet: a handler with a real bug in it still prints, or the
        # assertions below fail with no clue as to why.
        if not isinstance(sys.exc_info()[1], OSError):
            super().handle_error(request, client_address)


_srv = _LoopServer(("127.0.0.1", 0), _LoopHandler)
threading.Thread(target=_srv.serve_forever, daemon=True).start()
try:
    _root = "http://127.0.0.1:%d/" % _srv.server_address[1]
    _loop_scraper = WebScraper(base_url=_root, rate_limit=0)
    _loop_pages = _loop_scraper.scrape()
finally:
    _srv.shutdown()
    _srv.server_close()      # the listening socket, which shutdown() leaves open

assert [p["url"] for p in _loop_pages] == [_root], [p["url"] for p in _loop_pages]
assert _loop_scraper.total_errors == 0, _loop_scraper.total_errors
assert _LOOP_HITS == {"/": 1, "/robots.txt": 1, "/big.jpg": 1, "/endless": 1}, _LOOP_HITS
# Cut short, never downloaded whole. The bound is loose on purpose: the kernel
# socket buffer decides how much the server lands before the client's close()
# takes effect, so an exact figure would be a flakier check, not a stricter one.
assert _LOOP_SENT["jpeg"] < _JPEG_BYTES, _LOOP_SENT["jpeg"]
# ...and the page cap held on a body that never ends. The server stops itself at
# twice the cap, so a cap that stopped working fails here instead of hanging CI.
# The counter trails the socket by up to one write — it is bumped after the
# write returns, and the reader can be served from the buffer in between — so
# the floor allows a chunk. Without that slack the line passes only because the
# cap happens to be an exact multiple of the chunk (20 MB / 64 KB = 320), and
# would start failing at random the day the cap became an odd number.
assert (_crawl_MAX_PAGE_BYTES - _LOOP_CHUNK <= _LOOP_SENT["endless"]
        < 2 * _crawl_MAX_PAGE_BYTES), _LOOP_SENT["endless"]

# An image body is capped while it downloads, not after: with no Content-Length
# header the remote server would otherwise decide this run's memory use
class _BigImageResponse:
    status_code = 200
    headers = {"content-type": "image/png"}   # no content-length

    def close(self):
        pass

    def iter_content(self, n):
        sent = 0
        while sent < 20 * 1024 * 1024:
            sent += n
            yield b"x" * n


class _BigImageSession:
    headers = {}

    def get(self, *a, **kw):
        return _BigImageResponse()


import tempfile as _tf

from extract import ImageHandler

assert ImageHandler(_BigImageSession(), _tf.mkdtemp()).fetch(
    "https://example.com/big.png", "https://example.com/"
) is None

# Sparse stop happens during the crawl, not after it: a JS-rendered site builds
# no document, so crawling the rest of its tree at 1 req/s buys nothing.
# max_pages bounds the stub's link tree, so losing the stop fails the check
# rather than hanging it.
def _sparse_scraper(js_rendered, block=None):
    x = _no_net_scraper(None)
    x.js_rendered = js_rendered
    x.max_pages = 40
    n = [0]

    def tiny(url, depth):
        n[0] += 1
        return {"url": url, "title": "t", "binary_references": [],
                "content": [block or {"type": "paragraph", "text": "tiny"}],
                "links": [f"https://example.com/p{n[0]}a", f"https://example.com/p{n[0]}b"]}

    x._fetch_page = tiny
    return x


assert len(_sparse_scraper(True).scrape()) == SPARSE_PROBE_PAGES
# ...but the probe alone may not end a crawl. A nav-only landing page extracts to
# nothing too — correctly, nav is noise — so a static site runs to completion and
# the post-crawl average decides, exactly as before the early stop existed.
assert len(_sparse_scraper(False).scrape()) == 40, "static site must not be cut short"
# ...and a picture counts here exactly as it counts after the crawl. A photo
# archive reads as JavaScript-rendered — it has almost no visible text — so the
# text-only probe stopped it at ten pages, and main.py, which does count
# pictures, then published those ten as a complete archive of a whole gallery.
_PIC = {"type": "image", "src": "https://example.com/a.png", "alt": "",
        "mime_type": None, "path": None}
assert len(_sparse_scraper(True, _PIC).scrape()) == 40,     "a picture-only site must not be cut short"

# Word table cells: markdown markers format through _inline, and the header
# row stays bold. Needs python-docx; skipped with a message when missing.
try:
    import base64
    import tempfile
    from docx import Document
    from docx.shared import Cm
    from build_docx import MAX_IMAGE_HEIGHT, build_docx
    p2 = {
        "url": "https://example.com/", "title": "T", "breadcrumb": [],
        "content": [{"type": "table", "data": [["**H**", "x"], ["`c`", "plain"]]}],
        "links": [], "binary_references": [], "depth": 0,
    }
    d2 = tempfile.mkdtemp()
    dx2 = os.path.join(d2, "t.docx")
    build_docx([p2], dx2, "https://example.com/")
    t2 = Document(dx2).tables[0]
    runs = lambda c: [r for r in t2.cell(*c).paragraphs[0].runs if r.text]
    hdr = runs((0, 0))
    assert [r.text for r in hdr] == ["H"] and all(r.bold for r in hdr), [(r.text, r.bold) for r in hdr]
    code = runs((1, 0))
    assert [r.text for r in code] == ["c"] and code[0].font.name == "Courier New", [(r.text, r.font.name) for r in code]

    # Nesting renders through Word's built-in level styles, and an unreachable
    # domain is recorded in the archive rather than only on the console
    p3 = {
        "url": "https://example.com/", "title": "T", "breadcrumb": [],
        "content": [
            {"type": "list", "ordered": True, "level": 0, "start": 0, "items": ["Parent"]},
            {"type": "list", "ordered": False, "level": 1, "start": 0, "items": ["alpha"]},
            {"type": "list", "ordered": True, "level": 0, "start": 1, "items": ["two"]},
        ],
        "links": [], "binary_references": [], "depth": 0,
    }
    dx3 = os.path.join(d2, "t3.docx")
    build_docx([p3], dx3, "https://example.com/", blocked_domains=["dead.example.com"])
    d3 = Document(dx3)
    styled = [(p.style.name, p.text) for p in d3.paragraphs if "List" in p.style.name]
    assert styled == [("List Number", "Parent"), ("List Bullet 2", "alpha"),
                      ("List Number", "two"), ("List Bullet", "dead.example.com")], styled
    assert any(p.text == "Domains not reachable during extraction" for p in d3.paragraphs)

    # A link inside emphasis reaches Word as its label, never as raw brackets:
    # the inline alternation matches "**[y](url)**" as one bold part, so the
    # link branch never sees it. Word keeps label text only, by design.
    p5 = {
        "url": "https://example.com/", "title": "T", "breadcrumb": [],
        "content": [{"type": "paragraph",
                     "text": "see **[y](https://e.com/a)** and *[z](https://e.com/b)*"}],
        "links": [], "binary_references": [], "depth": 0,
    }
    dx5 = os.path.join(d2, "t5.docx")
    build_docx([p5], dx5, "https://example.com/")
    para = next(p for p in Document(dx5).paragraphs if p.text.startswith("see "))
    assert para.text == "see y and z", repr(para.text)
    assert [(r.text, bool(r.bold), bool(r.italic)) for r in para.runs if r.text.strip()] == [
        ("see ", False, False), ("y", True, False), (" and ", False, False), ("z", False, True),
    ], [(r.text, r.bold, r.italic) for r in para.runs]

    # main.py strips the base64 out of pages.json — the file would be enormous —
    # and keeps `path`. Both builders read `data` only, so a rebuild from that
    # file turned every picture into a placeholder while the image itself sat
    # unread in the work directory. Verified live: 1 image in, 0 embedded.
    import struct as _struct
    import zlib as _zlib

    def _png_bytes(w, h):
        def _c(tag, data):
            return (_struct.pack(">I", len(data)) + tag + data
                    + _struct.pack(">I", _zlib.crc32(tag + data)))
        return (bytes([137, 80, 78, 71, 13, 10, 26, 10])
                + _c(b"IHDR", _struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
                + _c(b"IEND", b""))

    _img_path = os.path.join(d2, "on-disk.png")
    with open(_img_path, "wb") as _f:
        _f.write(_png_bytes(120, 90))
    _slim = {
        "url": "https://example.com/", "title": "T", "breadcrumb": [],
        "content": [{"type": "image", "src": "https://example.com/a.png",
                     "alt": "pic", "mime_type": "image/png", "path": _img_path}],
        "links": [], "binary_references": [], "depth": 0,
    }
    dx6 = os.path.join(d2, "t6.docx")
    build_docx([_slim], dx6, "https://example.com/")
    _rebuilt = Document(dx6)
    assert len(_rebuilt.inline_shapes) == 1, len(_rebuilt.inline_shapes)
    assert not [x for x in _rebuilt.paragraphs if "[Image" in x.text], "placeholder, not picture"
    # ...and a path that no longer exists still degrades to the placeholder
    _gone = {**_slim, "content": [{**_slim["content"][0], "path": _img_path + ".gone"}]}
    dx7 = os.path.join(d2, "t7.docx")
    build_docx([_gone], dx7, "https://example.com/")
    assert [x.text for x in Document(dx7).paragraphs if "[Image" in x.text] == [
        "[Image: pic — https://example.com/a.png]"], "a missing file must not crash the build"

    # The disk read honours the same 5 MB ceiling the download does. A pages
    # file is an input like any other, and nothing ImageHandler wrote can be
    # over the cap, so only a hand-edited or crafted one reaches this.
    from common import MAX_IMAGE_BYTES as _CAP
    from common import image_bytes as _ib

    _huge = os.path.join(d2, "huge.png")
    with open(_huge, "wb") as _f:
        _f.write(b"x" * (_CAP + 1))
    assert _ib({"path": _huge}) is None, "the disk read must honour the image cap"
    assert _ib({"path": _img_path}) is not None
    assert _ib({}) is None and _ib({"path": _huge + ".gone"}) is None

    # A tall screenshot is scaled by height, not left running off the page.
    # Header-only PNG: python-docx reads the size from IHDR and decodes nothing,
    # so the check needs no image library.
    import struct
    import zlib

    def _png_chunk(tag, data):
        return (struct.pack(">I", len(data)) + tag + data
                + struct.pack(">I", zlib.crc32(tag + data)))

    png_signature = bytes([137, 80, 78, 71, 13, 10, 26, 10])
    tall_png = (png_signature
                + _png_chunk(b"IHDR", struct.pack(">IIBBBBB", 600, 4000, 8, 2, 0, 0, 0))
                + _png_chunk(b"IEND", b""))
    # Base64 with no path: nothing produces this shape any more, and a pages
    # file written before that change still holds it. This check is what keeps
    # image_bytes' legacy branch honest.
    p4 = {
        "url": "https://example.com/", "title": "T", "breadcrumb": [],
        "content": [{"type": "image", "alt": "tall", "src": "s", "mime_type": "image/png",
                     "data": base64.b64encode(tall_png).decode("ascii")}],
        "links": [], "binary_references": [], "depth": 0,
    }
    dx4 = os.path.join(d2, "t4.docx")
    build_docx([p4], dx4, "https://example.com/")
    shape = Document(dx4).inline_shapes[0]
    assert shape.height <= MAX_IMAGE_HEIGHT, shape.height / 360000
    assert shape.width < Cm(5), shape.width / 360000   # scaled down, not squashed
except ImportError:
    print("python-docx not installed - docx table check skipped")

# Markdown: nesting indents by four spaces (enough to clear a "1. " marker) and
# an ordered list keeps counting across the sublist
from build_md import _render, build_md
assert _render({"type": "list", "ordered": False, "level": 1, "start": 0,
                "items": ["alpha"]})[0] == "    - alpha"
assert _render({"type": "list", "ordered": True, "level": 0, "start": 1,
                "items": ["two"]})[0] == "2. two"

# A <br> reaches the builder as a newline, and a raw newline in a list item ends
# the item — the next line, opening with "-", became a list of its own
assert _render({"type": "list", "ordered": False, "level": 0, "start": 0,
                "items": ["first\n- still the same item"]})[0] == \
    "- first<br>- still the same item"
assert _render({"type": "paragraph", "text": "line\n# not a heading"})[0] == \
    "line<br># not a heading"

# Markdown: a page title is written as "## {title}", so the page's own <h2> is
# the level below it. Offsetting by two skipped that level outright and flattened
# h4, h5 and h6 into one — the deepest three headings of an API reference all
# came out "######". Word and PDF both offset by one, and this was the third
# builder disagreeing with them.
assert [_render({"type": "heading", "level": n, "text": "X"})[0]
        for n in (2, 3, 4, 5, 6)] == ["### X", "#### X", "##### X",
                                      "###### X", "###### X"]

# PDF: the extractor's markdown markers format in list items, table cells and
# callouts too, not only in paragraphs — Word already formats all four
from build_pdf import _block_html

assert _block_html({"type": "list", "ordered": False, "level": 0, "start": 0,
                    "items": ["see **bold**"]}) == "<ul><li>see <strong>bold</strong></li></ul>"
assert "<strong>H</strong>" in _block_html({"type": "table", "data": [["**H**", "x"]]})
assert _block_html({"type": "callout", "text": "note **now**"}) ==     "<blockquote>note <strong>now</strong></blockquote>"

# Markdown: a literal | in a cell would add a column and break the whole table
assert _render({"type": "table", "data": [["A|B", "c"]]})[0] == r"| A\|B | c |"
assert _render({"type": "table", "data": [["one\ntwo", "c"]]})[0] == "| one<br>two | c |"

# ...and the widest row sets the width, not the header. A colspan header extracts
# as one cell above two-cell body rows, and truncating to the header's width
# dropped "400" from the document with nothing said about it. Word and PDF kept it.
_ragged = _render({"type": "table", "data": [["Spec"], ["Torque", "400"]]})
assert _ragged[:3] == ["| Spec |  |", "| --- | --- |", "| Torque | 400 |"], _ragged

# A table inside a table is a block of its own: left in place, its rows were
# emitted twice and its text glued into the enclosing cell ("outerinner")
NESTED_TBL = ('<html><head><title>T</title></head><body><main>' + PAD +
              '<table><tr><td>outer<table><tr><td>inner</td></tr></table>'
              '</td></tr></table></main></body></html>')
nt = [b for b in PageExtractor().extract(NESTED_TBL, "https://example.com/")["content"]
      if b["type"] == "table"]
assert nt == [{"type": "table", "data": [["outer"]]},
              {"type": "table", "data": [["inner"]]}], nt

# Table cells keep the inline markers the builders already render; every builder
# formatted them, only the extractor threw them away
CELL_FMT = ('<html><head><title>T</title></head><body><main>' + PAD +
            '<table><tr><th>H</th></tr>'
            '<tr><td><strong>b</strong> and <code>c</code></td></tr>'
            '</table></main></body></html>')
cf = [b for b in PageExtractor().extract(CELL_FMT, "https://example.com/")["content"]
      if b["type"] == "table"]
assert cf == [{"type": "table", "data": [["H"], ["**b** and `c`"]]}], cf

# Merged cells: a colspan header is one cell above two body cells, and a rowspan
# label is written once and missing from the rows below it, so reading each <tr>
# in cell order shifted every value to their right one column left. Measured:
# "Val" landed over "b" and "c" under the "Spec" heading — wrong values under
# real headings, with nothing in the document saying so.
SPANS = ('<html><head><title>T</title></head><body><main>' + PAD +
         '<table>'
         '<tr><th colspan="2">Spec</th><th>Val</th></tr>'
         '<tr><td rowspan="2">A</td><td>b</td><td>1</td></tr>'
         '<tr><td>c</td><td>2</td></tr>'
         '</table></main></body></html>')
sp = [b for b in PageExtractor().extract(SPANS, "https://example.com/")["content"]
      if b["type"] == "table"]
assert sp == [{"type": "table", "data": [["Spec", "Spec", "Val"],
                                         ["A", "b", "1"],
                                         ["A", "c", "2"]]}], sp
# ...and the grid reaches the built document, not only the block
assert _render(sp[0])[:2] == ["| Spec | Spec | Val |", "| --- | --- | --- |"], _render(sp[0])

# A colspan is page-authored, so it is untrusted: one cell claiming 99999 columns
# must not decide how wide the row — and the Word table built from it — becomes
from extract import MAX_TABLE_COLUMNS

_wide = [b for b in PageExtractor().extract(
    '<html><head><title>T</title></head><body><main>' + PAD +
    '<table><tr><td colspan="99999">x</td></tr></table></main></body></html>',
    "https://example.com/")["content"] if b["type"] == "table"]
assert len(_wide[0]["data"][0]) == MAX_TABLE_COLUMNS, len(_wide[0]["data"][0])

# A <caption> is the table's title, and this branch reads <tr> and then returns,
# so "Table 3: Torque in Nm" reached no block and no builder at all. It leaves
# the tree the way the nested tables do, or _table_rows walks over it.
_cap = PageExtractor().extract(
    '<html><head><title>T</title></head><body><main>' + PAD +
    '<table><caption>Table 3: Torque in <b>Nm</b></caption>'
    '<tr><th>Bolt</th></tr><tr><td>M8</td></tr></table></main></body></html>',
    "https://example.com/")["content"]
assert _cap[-2:] == [
    {"type": "paragraph", "text": "Table 3: Torque in **Nm**"},
    {"type": "table", "data": [["Bolt"], ["M8"]]},
], _cap[-2:]

# An empty <pre> is furniture the page never showed: it printed as an empty grey
# box in Word and PDF and as an empty fence in Markdown
assert not [b for b in PageExtractor().extract(
    '<html><head><title>T</title></head><body><main>' + PAD +
    '<pre></pre><pre>   </pre></main></body></html>',
    "https://example.com/")["content"] if b["type"] == "code"]

# Lazy-loaded images: a script fills src in later, so the real address sits in
# data-src or srcset. Without the fallback the picture left no block at all.
LAZY = ('<html><head><title>T</title></head><body><main>' + PAD +
        '<img data-src="/a.png" alt="A">'
        '<img data-original="/b.png" alt="B">'
        '<img srcset="/c-480.png 480w, /c-960.png 960w" alt="C">'
        '<img alt="none">'
        '</main></body></html>')
lz = [(b["src"], b["alt"]) for b in PageExtractor().extract(LAZY, "https://example.com/")["content"]
      if b["type"] == "image"]
assert lz == [("https://example.com/a.png", "A"), ("https://example.com/b.png", "B"),
              ("https://example.com/c-480.png", "C")], lz

# Links survive into the block text as markdown. Plain label whenever the
# markdown form would be ambiguous or the target is not a page.
LINKS = ('<html><head><title>T</title></head><body><main>' + PAD +
         '<p><a href="mailto:x@example.com">mail</a></p>'
         '<p><a href="/rel">rel</a></p>'
         '<p><a href="/x">label [with] brackets</a></p>'
         '<p><a href="/wiki/Foo_(bar)">parens</a></p>'
         '<p><a href="/y"></a></p>'
         '</main></body></html>')
lk = [b["text"] for b in PageExtractor().extract(LINKS, "https://example.com/docs/")["content"]
      if b["type"] == "paragraph"]
assert lk[1:] == ["mail", "[rel](https://example.com/rel)",
                  "label [with] brackets", "parens"], lk

# robots.txt is a per-host file: one host's rules must never govern another, and
# each host is fetched once
class _RobotsSession:
    def __init__(self, bodies):
        self.bodies, self.calls = bodies, []

    def get(self, url, **kw):
        self.calls.append(url)
        body = self.bodies.get(url)

        # iter_content and close, because robots.txt is read to a cap like every
        # other untrusted body — a real one is a few KB, and an endless one used
        # to decide how much memory the run used
        class R:
            status_code = 200 if body is not None else 404
            text = body or ""

            def iter_content(self, n):
                yield self.text.encode()

            def close(self):
                pass
        return R()


rs = _RobotsSession({
    "https://a.example.com/robots.txt": "User-agent: *\nDisallow: /private\n",
    "https://b.example.com/robots.txt": "User-agent: *\nDisallow: /secret\n",
})
rx = WebScraper.__new__(WebScraper)
rx.session = rs
rx._robots_by_host = {}
assert rx._allowed_by_robots("https://a.example.com/ok") is True
assert rx._allowed_by_robots("https://a.example.com/private/x") is False
assert rx._allowed_by_robots("https://b.example.com/private/x") is True, "A's rules governed B"
assert rx._allowed_by_robots("https://b.example.com/secret/x") is False
assert rx._allowed_by_robots("https://c.example.com/any") is True   # no robots.txt → allowed
assert len(rs.calls) == 3, rs.calls   # one fetch per host, then cached

# Every sitemap robots.txt declares, not only the last: sites split them by
# section, and keeping one lost all the other pages
from crawl import _sitemap_lines

assert _sitemap_lines(
    "User-agent: *\nSitemap: https://e.com/sitemap-posts.xml\n"
    "Disallow: /x\nsitemap: https://e.com/sitemap-pages.xml\n"
) == ["https://e.com/sitemap-posts.xml", "https://e.com/sitemap-pages.xml"]
assert _sitemap_lines("User-agent: *\nDisallow: /\n") == []

# A blank query value marks a distinct page, so the two must not share one
# visited key — one of them would never be fetched
assert (WebScraper._normalize(dummy, "https://example.com/a?x=&y=1")
        != WebScraper._normalize(dummy, "https://example.com/a?y=1"))

# Pages withheld by robots.txt are recorded, so the archive shows the gap
import urllib.robotparser

_rp = urllib.robotparser.RobotFileParser()
_rp.parse(["User-agent: *", "Disallow: /private"])
rex = _no_net_scraper(None)
rex._robots_by_host = {"example.com": _rp}
rex.scrape(sitemap_urls=["https://example.com/private/p", "https://example.com/ok"])
assert rex.robots_excluded == ["https://example.com/private/p"], rex.robots_excluded

_md_path = os.path.join(_tf.mkdtemp(), "out.md")
build_md([{"url": "https://example.com/", "title": "T", "breadcrumb": [],
           "content": [], "binary_references": [], "depth": 0}],
         _md_path, "https://example.com/", robots_excluded=rex.robots_excluded)
_md_text = open(_md_path, encoding="utf-8").read()
assert "Pages excluded by robots.txt" in _md_text
assert "https://example.com/private/p" in _md_text

# ...but the listing is capped: a robots.txt disallowing a whole branch can
# exclude thousands of URLs on an unlimited-depth crawl, and the reference
# section prints every one it is handed
from main import cap_listing

_capped = cap_listing([f"https://example.com/p{i}" for i in range(60)])
assert len(_capped) == 51, len(_capped)
assert _capped[-1] == "10 more not listed", _capped[-1]
# build_md writes each entry as "- {entry}", so a tail line opening with a second
# list marker would nest into a sublist and lose its own text
assert not _capped[-1].startswith(("-", "+", "*", ">")), _capped[-1]
# no "+ N more" line when nothing was dropped, at the boundary or below it
assert cap_listing([f"u{i}" for i in range(50)]) == [f"u{i}" for i in range(50)]
assert cap_listing(["a", "b"]) == ["a", "b"]
assert cap_listing([f"u{i}" for i in range(51)])[-1] == "1 more not listed"

# Markdown: a code sample that itself shows a fence must not end its own block
assert _render({"type": "code", "language": "", "text": "```\nx\n```"})[0] == "````"
assert _render({"type": "code", "language": "py", "text": "print(1)"})[0] == "```py"

# Markdown: two pages titled the same get distinct anchors, GitHub's own rule —
# otherwise both contents links jump to the first page
from build_md import _unique_anchor

_seen: dict = {}
assert [_unique_anchor("intro", _seen) for _ in range(3)] == ["intro", "intro-1", "intro-2"]

# PDF renders the markdown link as a real anchor; Word renders the label only,
# because python-docx has no hyperlink API and raw brackets would be worse
assert _block_html({"type": "paragraph", "text": "see [the **big** guide](https://e.com/a)"}) == (
    '<p>see <a href="https://e.com/a">the <strong>big</strong> guide</a></p>')

# The PDF builder reads the picture from disk too, and re-encodes it into the
# data URI weasyprint embeds — the same rebuild that had lost it in Word
_pdf_img = os.path.join(_tf.mkdtemp(), "p.png")
with open(_pdf_img, "wb") as _f:
    _f.write(b"picture payload the builder never has to parse")
_html_img = _block_html({"type": "image", "alt": "pic", "src": "https://e.com/a.png",
                         "mime_type": "image/png", "path": _pdf_img})
assert _html_img.startswith('<img src="data:image/png;base64,'), _html_img
assert _block_html({"type": "image", "alt": "pic", "src": "https://e.com/a.png",
                    "path": _pdf_img + ".gone"}).startswith('<p class="img-placeholder"')

# PDF needs an engine; with neither present the build must fail loudly, which is
# the path a machine without LibreOffice actually takes
import build_pdf as _bp

try:
    import weasyprint  # noqa: F401
except ImportError:
    _bp._find_libreoffice = lambda: None
    try:
        _bp.build_pdf([], os.path.join(_tf.mkdtemp(), "x.pdf"), "https://example.com/")
        raise AssertionError("build_pdf must exit non-zero when no engine is available")
    except SystemExit as _e:
        assert _e.code == 1, _e.code

# ...and a weasyprint that is installed but unusable must fail the same way. On
# Windows `pip install weasyprint` succeeds and the import then raises OSError
# from cffi, because the GTK libraries it binds to are not pip-installable there.
# That is the exact path SKILL.md sends a user down when LibreOffice is absent,
# and `except ImportError` let it escape as a traceback.
_fake_dir = _tf.mkdtemp()
with open(os.path.join(_fake_dir, "weasyprint.py"), "w", encoding="utf-8") as _f:
    _f.write("raise OSError(\"cannot load library 'libgobject-2.0-0'\")\n")


def _forget_weasyprint():
    """Drop the package and its submodules, so the next import is a fresh one.

    In CI the real weasyprint is installed and already imported by the check
    above; leaving stale submodules behind would make the re-import below read
    half of one package and half of the other.
    """
    for _m in [m for m in sys.modules if m == "weasyprint" or m.startswith("weasyprint.")]:
        del sys.modules[_m]


sys.path.insert(0, _fake_dir)
_forget_weasyprint()
_real_lo, _bp._find_libreoffice = _bp._find_libreoffice, lambda: None
try:
    assert _bp._import_weasyprint().startswith("OSError:"), _bp._import_weasyprint()
    try:
        _bp.build_pdf([], os.path.join(_tf.mkdtemp(), "x.pdf"), "https://example.com/")
        raise AssertionError("an unusable weasyprint must exit 1, not raise OSError")
    except SystemExit as _e:
        assert _e.code == 1, _e.code
finally:
    _bp._find_libreoffice = _real_lo
    sys.path.remove(_fake_dir)
    _forget_weasyprint()

# ...and when an engine IS present the PDF path runs end to end. Until this check
# it was the only never-executed code in the skill: no LibreOffice and no
# weasyprint on the maintainer's machine, and CI installed neither.
try:
    import weasyprint as _wp
except ImportError:
    print("weasyprint not installed - PDF end-to-end check skipped")
else:
    _pdf_pages = [
        {"url": "https://example.com/a", "title": "First", "breadcrumb": ["Docs"],
         "depth": 0, "binary_references": [],
         "content": [{"type": "heading", "level": 2, "text": "Heading"},
                     {"type": "paragraph", "text": "Body with **bold** and a "
                                                   "[link](https://e.com/x)."},
                     {"type": "table", "data": [["Spec"], ["Torque", "400"]]}]},
        {"url": "https://example.com/b", "title": "Second", "breadcrumb": [],
         "depth": 1, "binary_references": [],
         "content": [{"type": "list", "ordered": True, "level": 0, "start": 0,
                      "items": ["one"]}]},
    ]
    # Four pages: cover, contents, and one per page. Three would mean the cover
    # and the contents list shared a page, which is what h1:first-of-type did.
    _doc = _wp.HTML(string=_bp._build_html(_pdf_pages, "https://example.com/")).render()
    assert len(_doc.pages) == 4, len(_doc.pages)

    # Force the weasyprint strategy: LibreOffice is the preferred engine when
    # installed, and it is the one branch this check cannot reach.
    _real_lo, _bp._find_libreoffice = _bp._find_libreoffice, lambda: None
    _pdf_out = os.path.join(_tf.mkdtemp(), "site.pdf")
    _bp.build_pdf(_pdf_pages, _pdf_out, "https://example.com/")
    _bp._find_libreoffice = _real_lo
    with open(_pdf_out, "rb") as _f:
        _pdf_bytes = _f.read()
    assert _pdf_bytes.startswith(b"%PDF-"), _pdf_bytes[:20]
    assert len(_pdf_bytes) > 2000, len(_pdf_bytes)

# Inline markers and links survive wherever the text sits: straight in a
# container with no <p> wrapper, and inside a block child of a list item. While
# the extractor had one inline loop per caller, only the <p> path applied them.
INL = ('<html><head><title>T</title></head><body><main>' + PAD +
       '<div>plain <strong>bold</strong> <em>it</em> <code>c</code> '
       '<a href="/l">link</a></div>'
       '<ul><li><div><strong>deep</strong> <a href="/d">dl</a></div></li></ul>'
       '</main></body></html>')
_inl = PageExtractor().extract(INL, "https://example.com/")["content"]
assert _inl[1]["text"] == "plain **bold** *it* `c` [link](https://example.com/l)", _inl[1]
assert _inl[2]["items"] == ["**deep** [dl](https://example.com/d)"], _inl[2]

# An icon font is an empty inline element. Five <i class="icon-star"></i> in a
# row are five empty italics, and marking them printed "** ** ** ** **" where
# the page showed a star rating (live-verified on books.toscrape.com).
ICONS = ('<html><head><title>T</title></head><body><main>' + PAD +
         '<p><i class="icon-star"></i><i class="icon-star"></i>'
         '<strong> </strong><code></code>Three stars</p>'
         '</main></body></html>')
_ico = [b["text"] for b in PageExtractor().extract(ICONS, "https://example.com/")["content"]]
assert _ico[1] == "Three stars", repr(_ico[1])

# An inline wrapper adds no space of its own — <sup>, <kbd> and friends render
# tight against their neighbours, and a space between every inline element broke
# "x2y" into "x 2 y". A marker also has to hug its text: "** bold **" is literal
# asterisks in Markdown, not emphasis, and pretty-printed HTML puts the newlines
# inside the tag.
TIGHT = ('<html><head><title>T</title></head><body><main>' + PAD +
         '<p>x<sup>2</sup>y and <abbr>WHO</abbr>-led</p>'
         """<p>a<strong>
  bold
</strong>b</p>"""
         '</main></body></html>')
_tight = [b["text"] for b in PageExtractor().extract(TIGHT, "https://example.com/")["content"]]
assert _tight[1] == "x2y and WHO-led", repr(_tight[1])
assert _tight[2] == "a **bold** b", repr(_tight[2])

# HTTP 429 is the site asking for a slower pace, not a dead page: wait the
# interval it names and retry once, rather than spending the error budget.
# crawl.time is stubbed so the checks below pin the wait instead of taking it.
import types

import crawl as _crawl

_slept: list = []
_crawl.time = types.SimpleNamespace(sleep=_slept.append)


class _ThrottleSession:
    """429 first, then 200 — the throttled page that comes good on retry."""

    def __init__(self, always=False):
        self.always, self.calls = always, 0

    def get(self, url, **kw):
        self.calls += 1
        throttled = self.always or self.calls == 1

        class R:
            status_code = 429 if throttled else 200
            headers = {"Retry-After": "0", "content-type": "text/html"}
            text = "<html><body><main><p>ok</p></main></body></html>"
            content = text.encode()
            encoding = "utf-8"
            url = "https://example.com/"

            def close(self):
                pass

            def iter_content(self, n):
                yield self.content

            def raise_for_status(self):
                if self.status_code != 200:
                    raise requests.exceptions.HTTPError(str(self.status_code))
        return R()


def _throttled_scraper(always=False):
    x = _no_net_scraper(None)
    del x._fetch_page
    x.session = _ThrottleSession(always)
    x.extractor = PageExtractor()
    return x


t = _throttled_scraper()
assert len(t.scrape()) == 1, "the retried page must be kept"
assert t.session.calls == 2, t.session.calls
assert t.total_errors == 0, t.total_errors
# a Retry-After of 0 is still floored: the retry never re-asks with no gap at all
assert _slept == [1], _slept

# ...and the retry is strictly one. A host that answers 429 for ever must not
# recurse: the second 429 is a page error like any other. Without the retry=False
# argument this recursed ~1000 deep, sleeping between each level, and the
# RecursionError was swallowed by the blanket except in _fetch_page.
_slept.clear()
t = _throttled_scraper(always=True)
assert t.scrape() == []
assert t.session.calls == 2, t.session.calls
assert t.total_errors == 1, t.total_errors
assert _slept == [1], _slept

# A Crawl-delay is part of robots.txt, and the skill claims robots.txt is always
# respected. Only Disallow was read, so a host asking 10s between requests got 1.
from crawl import CRAWL_DELAY_MAX, ROBOTS_AGENT, _asked_delay


def _robots(*lines):
    r = urllib.robotparser.RobotFileParser()
    r.parse(["User-agent: *", *lines])
    return r


assert _asked_delay(None) == 0.0
assert _asked_delay(_robots("Disallow: /x")) == 0.0
assert _asked_delay(_robots("Crawl-delay: 4")) == 4.0
# Request-rate is the standardised sibling. Both are floors on the gap, so the
# slower one governs — including when Request-rate would permit a faster pace.
assert _asked_delay(_robots("Request-rate: 1/5")) == 5.0
assert _asked_delay(_robots("Crawl-delay: 4", "Request-rate: 1/9")) == 9.0
assert _asked_delay(_robots("Crawl-delay: 9", "Request-rate: 10/1")) == 9.0

# Adopted at the START host, whose robots.txt the analyzer already fetched and
# handed over: _allowed_by_robots never fetches for that host, so a hook in the
# fetch path alone would have missed the commonest case of all.
_ws = WebScraper(base_url="https://example.com/", rate_limit=1.0,
                 robots=_robots("Crawl-delay: 4"))
assert _ws.rate_limit == 4.0, _ws.rate_limit
assert _ws.image_handler.rate_limit == 4.0, "the image downloader shares the pace"

# ...clamped, so a robots.txt asking for an hour cannot hang a run
assert WebScraper(base_url="https://example.com/", rate_limit=1.0, embed_images=False,
                  robots=_robots("Crawl-delay: 9999")).rate_limit == CRAWL_DELAY_MAX

# ...and only ever raised. A user asking to be slower than the site stays slower.
assert WebScraper(base_url="https://example.com/", rate_limit=10.0, embed_images=False,
                  robots=_robots("Crawl-delay: 4")).rate_limit == 10.0
# ...including slower than the cap. CRAWL_DELAY_MAX bounds what the SITE asks for,
# never the caller's own floor — capping both cut a --rate-limit of 60 down to 30.
assert WebScraper(base_url="https://example.com/", rate_limit=60.0, embed_images=False,
                  robots=_robots("Crawl-delay: 9999")).rate_limit == 60.0
# one function owns the rule, because main.py reports the pace and WebScraper
# enforces it — they disagreed, so Phase 1 printed a pace the crawl never used
from crawl import effective_delay

assert effective_delay(_robots("Crawl-delay: 9999"), 60.0) == 60.0
assert effective_delay(_robots("Crawl-delay: 9999"), 1.0) == CRAWL_DELAY_MAX
assert effective_delay(_robots("Disallow: /x"), 1.0) == 1.0

# Adopted at a host first seen MID-CRAWL too, through the fetch path. The class
# default is what lets a __new__-built scraper reach this at all.
_dx = WebScraper.__new__(WebScraper)
assert _dx.rate_limit == 0.0 and _dx.image_handler is None
_dx.session = _RobotsSession(
    {"https://slow.example.com/robots.txt": "User-agent: *\nCrawl-delay: 7\n"})
_dx._robots_by_host = {}
assert _dx._allowed_by_robots("https://slow.example.com/a") is True
assert _dx.rate_limit == 7.0, _dx.rate_limit

# And the crawl actually waits that long. crawl.time is still stubbed here, so
# this pins the pause instead of taking it.
_slept.clear()
_cd = _no_net_scraper(None)
_cd.rate_limit = 1.0
_cd.session = _RobotsSession(
    {"https://example.com/robots.txt": "User-agent: *\nCrawl-delay: 3\n"})
_cd._robots_by_host = {}          # force the fetch path, as a real run does
assert len(_cd.scrape()) == 1
assert _cd.rate_limit == 3.0, _cd.rate_limit
assert _slept == [3.0], _slept
_slept.clear()

from crawl import RETRY_AFTER_DEFAULT, RETRY_AFTER_MAX, RETRY_AFTER_MIN, _retry_after

assert _retry_after("3") == 3
assert _retry_after(None) == RETRY_AFTER_DEFAULT
assert _retry_after("whenever") == RETRY_AFTER_DEFAULT
assert _retry_after("9999") == RETRY_AFTER_MAX                # clamped, never hangs a run
# a 429 answer is untrusted input: 0, a negative, and a past HTTP-date (routine
# with any clock skew) must all still leave a gap before the retry
assert _retry_after("0") == RETRY_AFTER_MIN
assert _retry_after("-5") == RETRY_AFTER_MIN
assert _retry_after("Wed, 21 Oct 2015 07:28:00 GMT") == RETRY_AFTER_MIN

# The analyzer fetches the homepage before the crawler exists, so a site that
# throttles from its first request used to end the whole run at Phase 1 with
# "not reachable" — the one case the retry above is for.
_slept.clear()
_an = SiteAnalyzer.__new__(SiteAnalyzer)
_an.url = "https://example.com/"
_an.session = _ThrottleSession()
assert _an.analyze()["reachable"] is True, "a first-request 429 is throttling, not a dead site"
assert _an.session.calls > 1, _an.session.calls
assert _slept == [1], _slept

_crawl.time = time   # restore, so any later check sleeps for real

# The Chrome MCP branch is one script, and its dedup and its title fallback had
# no check at all. HTML with no <img> keeps this offline.
import json as _json

from chrome_extract import save_page

_pf = os.path.join(_tf.mkdtemp(), "pages.json")
_PAGE = ("<html><head><title>From title</title></head><body><main><h1>Real</h1>"
         "<p>Body text long enough to make this main a genuine content area.</p>"
         "</main></body></html>")
save_page(_pf, "https://example.com/a", "Argument title", _PAGE)
save_page(_pf, "https://example.com/a", "Argument title", _PAGE)   # same URL twice
save_page(_pf, "https://example.com/b", "Argument title",
          "<html><body><main><p>No heading, no title element.</p></main></body></html>")
_saved = _json.load(open(_pf, encoding="utf-8"))
assert [q["url"] for q in _saved] == ["https://example.com/a", "https://example.com/b"], _saved
assert _saved[0]["title"] == "Real", _saved[0]["title"]             # the page's own <h1> wins
assert _saved[1]["title"] == "Argument title", _saved[1]["title"]   # --title only when empty

# ...and its image downloads are paced. The browser fetched the page, but the
# pictures still come from the site, so a rendered page with 40 of them fired 40
# back-to-back requests while the crawler was pausing between pages. save_page
# imports the handler at call time, so the spy goes on the extract module.
import inspect

import extract as _ex

_ce_rates: list = []
_ce_internal: list = []


class _RateSpyHandler:
    def __init__(self, session, images_dir, rate_limit=0.0, allow_internal=False):
        _ce_rates.append(rate_limit)
        _ce_internal.append(allow_internal)

    def fetch(self, img_url, base_url):
        return None


_real_handler, _ex.ImageHandler = _ex.ImageHandler, _RateSpyHandler
save_page(os.path.join(_tf.mkdtemp(), "p.json"), "https://example.com/i", "T", _PAGE,
          rate_limit=2.5)
# ...and the page's own address decides the trust zone, so a rendered page saved
# from a public site cannot aim its image downloads at this machine's network,
# while one saved from a local server keeps its own pictures
save_page(os.path.join(_tf.mkdtemp(), "p.json"), "http://127.0.0.1:8000/i", "T", _PAGE)
_ex.ImageHandler = _real_handler
assert _ce_rates == [2.5, 1.0], _ce_rates
assert _ce_internal == [False, True], _ce_internal
# and the default is the crawler's own, not the handler's unthrottled 0.0
assert inspect.signature(save_page).parameters["rate_limit"].default == 1.0

# ...and saving is unconditional. It was gated behind a --save-page flag, so a
# call that forgot the flag exited 0, printed nothing and wrote nothing: a
# fifty-page extraction looked successful one page at a time and ended with an
# empty pages file. The directory is absent here too, because save_page owns
# making it rather than inheriting it from ImageHandler's images/ subdirectory.
import chrome_extract as _ce

_bare = os.path.join(_tf.mkdtemp(), "not-created-yet", "pages.json")
sys.argv = ["chrome_extract.py", "--url", "https://example.com/f", "--title", "T",
            "--html", _PAGE, "--pages-file", _bare]
_ce.main()
assert [q["url"] for q in _json.load(open(_bare, encoding="utf-8"))] == [
    "https://example.com/f"], _bare

# Never overwrite an earlier archive of the same site
from main import non_clobbering_path

_arch = os.path.join(_tf.mkdtemp(), "site.docx")
assert non_clobbering_path(_arch) == _arch
open(_arch, "w").close()
assert non_clobbering_path(_arch) == _arch[:-5] + "_v2.docx"
open(_arch[:-5] + "_v2.docx", "w").close()
assert non_clobbering_path(_arch) == _arch[:-5] + "_v3.docx"


# A --pages-file is read from disk and appended to one page at a time, so a
# half-written or hand-edited one is a real possibility. The wrong shape must fail
# with a line of its own, not an AttributeError from inside a builder.
import main as _main

_bad_dir = _tf.mkdtemp()
for _body in ('{"a": 1}', '["a", "b"]'):
    _bad = os.path.join(_bad_dir, "bad.json")
    open(_bad, "w", encoding="utf-8").write(_body)
    sys.argv = ["main.py", "--url", "https://e.com", "--pages-file", _bad,
                "--format", "md", "--output", os.path.join(_bad_dir, "o.md")]
    try:
        _main.main()
    except SystemExit as _e:
        assert _e.code == 1, _e.code
    else:
        raise AssertionError(f"a malformed pages file was accepted: {_body}")

# Path filters: archive one section of a large site. Exclude beats include, an
# empty include list means any path, and the domain rules still apply on top.
_ps = WebScraper(base_url="https://example.com/", rate_limit=0, embed_images=False,
                 include_paths=["/docs/"], exclude_paths=["/docs/legacy/"])
assert _ps._in_scope("https://example.com/docs/intro") is True
assert _ps._in_scope("https://example.com/blog/post") is False
assert _ps._in_scope("https://example.com/docs/legacy/v1") is False   # exclude wins
assert _ps._in_scope("https://other.example.org/docs/intro") is False  # domain still rules

_pe = WebScraper(base_url="https://example.com/", rate_limit=0, embed_images=False,
                 exclude_paths=["/private/"])
assert _pe._in_scope("https://example.com/anything") is True          # no include list
assert _pe._in_scope("https://example.com/private/x") is False

# No filters at all must behave exactly as before
_pn = WebScraper(base_url="https://example.com/", rate_limit=0, embed_images=False)
assert _pn._in_scope("https://example.com/whatever") is True

# The leading slash is optional, so one command works in Git Bash on Windows too
_pw = WebScraper(base_url="https://example.com/", rate_limit=0, embed_images=False,
                 include_paths=["docs/"], exclude_paths=["docs/legacy/"])
assert _pw.include_paths == ["/docs/"] and _pw.exclude_paths == ["/docs/legacy/"]
assert _pw._in_scope("https://example.com/docs/intro") is True
assert _pw._in_scope("https://example.com/docs/legacy/v1") is False

# A blank prefix is dropped, not kept as "/". SKILL.md hands the agent a command
# template, so an unset substitution arrives here as "" — and "/" as an exclude
# prefix would match every page and cut the archive to its start page in silence.
_pb = WebScraper(base_url="https://example.com/", rate_limit=0, embed_images=False,
                 include_paths=[""], exclude_paths=["", "  "])
assert _pb.include_paths == [] and _pb.exclude_paths == [], (_pb.include_paths, _pb.exclude_paths)
assert _pb._in_scope("https://example.com/anything") is True

# --rate-limit has to reach image downloads. They are most of the requests a
# picture-heavy site receives, and they used to fire back to back while only the
# page loop paused.
from extract import ImageHandler as _IH

_IMG_RATE = 0.02


class _ImgResp:
    status_code = 200
    headers = {"content-type": "image/png", "content-length": "3"}

    def iter_content(self, n):
        yield b"abc"

    def close(self):
        pass


class _ImgSession:
    headers: dict = {}

    def get(self, url, **kw):
        return _ImgResp()


_ih = _IH(_ImgSession(), _tf.mkdtemp(), rate_limit=_IMG_RATE)
_t0 = time.time()
for _n in range(5):
    _ih.fetch(f"/i{_n}.png", "https://example.com/")
_elapsed = time.time() - _t0
assert _elapsed >= 5 * _IMG_RATE, _elapsed

# A cache hit returns above the sleep: re-reading an image the site repeats on
# every page must not cost a pause it never spent a request on
_t0 = time.time()
_ih.fetch("/i0.png", "https://example.com/")
assert time.time() - _t0 < _IMG_RATE, "a cached image must not sleep"

# What a fetch hands back: where the file went, not the file. The cache holds
# one of these dicts per image URL for the whole crawl, so base64 in here was
# the picture kept a second time over.
_fetched = _ih.fetch("/i0.png", "https://example.com/")
assert set(_fetched) == {"mime_type", "path"}, _fetched
assert open(_fetched["path"], "rb").read() == b"abc"

# The crawl must start from where the site ANSWERS, not from what the user typed.
# The bug this pins lives in neither class — both are correct alone — but in the
# line of main() that wires them, so it is only reachable by running main().
import main as _main

_captured = {}


class _RedirectAnalyzer:
    # rate_limit, because Phase 1 paces its sitemap fetches now too
    def __init__(self, url, rate_limit=0.0):
        pass

    def analyze(self, read_sitemap=True):
        return {"reachable": True, "final_url": "https://angular.dev/",
                "js_rendered": False, "has_robots": False, "has_sitemap": False,
                "sitemap_page_count": 0, "sitemap_urls": [], "homepage_links": 7,
                "error": None}


class _CapturingScraper:
    def __init__(self, base_url, **kw):
        _captured["base_url"] = base_url
        self.blocked_domains, self.robots_excluded, self.total_errors = set(), [], 0

    def scrape(self, sitemap_urls=None):
        return [{"url": "https://angular.dev/", "title": "Home", "breadcrumb": [],
                 "depth": 0, "binary_references": [],
                 "content": [{"type": "paragraph", "text": "x" * 400}]}]


_crawl.SiteAnalyzer, _crawl.WebScraper = _RedirectAnalyzer, _CapturingScraper
_wd = _tf.mkdtemp()
sys.argv = ["main.py", "--url", "https://angular.io", "--format", "md",
            "--output", os.path.join(_wd, "out.md"), "--work-dir", _wd]
_main.main()
_crawl.SiteAnalyzer, _crawl.WebScraper = SiteAnalyzer, WebScraper   # restore
assert _captured["base_url"] == "https://angular.dev/", _captured

# A picture is content, and text_chars counts none. An image-only site — a photo
# archive, a slide portal — is thin on text by nature, not by failure, and used to
# be crawled in full at one request per second and then refused a document.
class _PlainAnalyzer(_RedirectAnalyzer):
    def analyze(self, read_sitemap=True):
        return {**_RedirectAnalyzer.analyze(self), "final_url": "https://example.com/"}


def _pages_scraper(contents):
    class S:
        def __init__(self, base_url, **kw):
            self.blocked_domains, self.robots_excluded, self.total_errors = set(), [], 0

        def scrape(self, sitemap_urls=None):
            return [{"url": f"https://example.com/{i}", "title": f"P{i}",
                     "breadcrumb": [], "depth": 0, "binary_references": [],
                     "content": c} for i, c in enumerate(contents)]
    return S


def _build_pages(contents, name):
    _crawl.SiteAnalyzer = _PlainAnalyzer
    _crawl.WebScraper = _pages_scraper(contents)
    d = _tf.mkdtemp()
    out = os.path.join(d, name)
    sys.argv = ["main.py", "--url", "https://example.com", "--format", "md",
                "--output", out, "--work-dir", d]
    _main.main()
    _crawl.SiteAnalyzer, _crawl.WebScraper = SiteAnalyzer, WebScraper
    return out


_IMG = {"type": "image", "src": "https://example.com/a.png", "alt": "",
        "mime_type": None, "path": None}
_TINY = {"type": "paragraph", "text": "hi"}

_gallery = _build_pages([[_IMG], [_IMG]], "gallery.md")
assert os.path.exists(_gallery), "an image-only site must still get a document"
assert "a.png" in open(_gallery, encoding="utf-8").read()

# ...and the safety net still holds for a site with neither text nor pictures
_empty = _build_pages([[_TINY]], "empty.md")
assert not os.path.exists(_empty), "a genuinely empty crawl must still build nothing"

# ...and one picture across a whole site is a logo, not a gallery. A JavaScript
# shell carrying its logo is exactly what the refusal exists for, and asking
# "any image at all" let it walk straight through.
_shell = _build_pages([[_IMG], [_TINY], [_TINY]], "shell.md")
assert not os.path.exists(_shell), "one logo must not defeat the sparse refusal"

# Phase 1 reports the pace the crawl will actually keep, not the one the site
# asked for: with the cap biting, those two numbers differ, and printing the
# asked-for one promised a 9999-second crawl that was never going to happen.
_delayed = urllib.robotparser.RobotFileParser()
_delayed.parse(["User-agent: *", "Crawl-delay: 9999"])


class _SlowAnalyzer(_PlainAnalyzer):
    def analyze(self, read_sitemap=True):
        return {**_PlainAnalyzer.analyze(self), "_rp": _delayed}


_crawl.SiteAnalyzer = _SlowAnalyzer
_crawl.WebScraper = _pages_scraper([[{"type": "paragraph", "text": "x" * 400}]])
_pd = _tf.mkdtemp()
sys.argv = ["main.py", "--url", "https://example.com", "--format", "md",
            "--output", os.path.join(_pd, "p.md"), "--work-dir", _pd]
_buf = io.StringIO()
with contextlib.redirect_stdout(_buf):
    _main.main()
_crawl.SiteAnalyzer, _crawl.WebScraper = SiteAnalyzer, WebScraper
_out = _buf.getvalue()
assert f"Pace    : {CRAWL_DELAY_MAX}s between requests" in _out, _out
assert "robots.txt asks for 9999s" in _out, _out

# <base href> governs every relative address on the page, the way a browser reads
# it. A framework app ships one by default, and rendered app HTML carries it into
# the Chrome MCP path, so joining on the page's own address sent every link and
# every picture to an address that does not exist.
_BASE_HTML = """<html><head><base href="https://example.com/v2/">
</head><body><main>
<p>See <a href="guide.html">guide</a> <img src="pic.png" alt="pic"></p>
<p>{pad}</p></main></body></html>""".format(pad="x" * 90)

_based = PageExtractor().extract(_BASE_HTML, "https://example.com/docs/intro.html")
assert _based["links"] == ["https://example.com/v2/guide.html"], _based["links"]
_bimg = [b["src"] for b in _based["content"] if b["type"] == "image"]
assert _bimg == ["https://example.com/v2/pic.png"], _bimg
# the page keeps its own address: only the join base moves, never the identity
# the visited set and the Source line are keyed on
assert _based["url"] == "https://example.com/docs/intro.html", _based["url"]
# the markdown link carried in the block text resolves through the same base
assert "(https://example.com/v2/guide.html)" in _based["content"][0]["text"]

# ...and with no <base> the page's own address is still the base
_plain = PageExtractor().extract(
    _BASE_HTML.replace('<base href="https://example.com/v2/">', ""),
    "https://example.com/docs/intro.html")
assert _plain["links"] == ["https://example.com/docs/guide.html"], _plain["links"]
_pimg = [b["src"] for b in _plain["content"] if b["type"] == "image"]
assert _pimg == ["https://example.com/docs/pic.png"], _pimg


# A server may answer from a host the user never agreed to. _in_scope gates the
# link, before the fetch, so a redirect walks straight past it and the off-domain
# page was archived under same-domain scope — an outbound-link redirector would
# have pulled arbitrary third-party pages into the archive.
_off = _no_net_scraper(None)
_off._fetch_page = lambda url, depth: {
    "url": ("https://other.example.net/landing" if url.endswith("/go") else url),
    "title": url, "content": [], "links": ["https://example.com/go"],
    "binary_references": [],
}
_pages = _off.scrape()
assert [q["url"] for q in _pages] == ["https://example.com/"], [q["url"] for q in _pages]
# a skip, not an error: a site redirecting many links off-domain must not spend
# the ten-consecutive-error budget
assert _off.consecutive_errors == 0, _off.consecutive_errors
assert _off.total_errors == 0, _off.total_errors

# ...while a redirect that stays inside the scope is still archived, at the
# address it answered from — the round-10 behaviour this must not undo
_in = _no_net_scraper(None)
_in._fetch_page = lambda url, depth: {
    "url": ("https://example.com/moved" if url.endswith("/go") else url),
    "title": url, "content": [], "links": ["https://example.com/go"],
    "binary_references": [],
}
assert [q["url"] for q in _in.scrape()] == [
    "https://example.com/", "https://example.com/moved"], _in.visited


# --depth and --max-pages are parsed at the boundary, not halfway through main().
# A bare int() raised a traceback after the work directory had already been
# created, and "full" is a likely thing to type — SKILL.md uses that word as the
# filename's depth label. argparse must answer with a usage line instead.
def _parses(*extra):
    sys.argv = ["main.py", "--url", "https://e.com", "--output", "o.md", *extra]
    return _main.parse_args()


for _bad in (["--depth", "full"], ["--max-pages", "lots"], ["--depth", "-1"],
             ["--max-pages", "0"]):          # 0 used to read as "no limit"
    try:
        # argparse prints its usage block to stderr on every rejection; four of
        # those would bury the rest of this run's output
        with contextlib.redirect_stderr(io.StringIO()):
            _parses(*_bad)
    except SystemExit:
        pass
    else:
        raise AssertionError(f"{_bad} was accepted")

assert _parses("--depth", "0").depth == 0                    # start page only
assert _parses("--depth", "unlimited").depth is None
assert _parses().depth is None and _parses().max_pages is None
assert _parses("--max-pages", "5").max_pages == 5


# ...but the start URL is archived wherever it lands. It is what the crawl
# follows links from, and --include-path deliberately does not gate it, so
# applying the scope check to its redirect emptied the whole run.
_start = _no_net_scraper(None)
_start.base_url = "https://site.com/"
_start.base_domain = "site.com"
_start._robots_by_host = {"site.com": None}
_start.include_paths = ("/docs/",)          # the start page sits outside this
_start._fetch_page = lambda url, depth: {
    "url": ("https://site.com/en/" if url == "https://site.com/" else url),
    "title": url, "content": [], "links": ["https://site.com/docs/a"],
    "binary_references": [],
}
assert [q["url"] for q in _start.scrape()] == [
    "https://site.com/en/", "https://site.com/docs/a"], _start.visited

# robots.txt is judged on the address that answered too, and exempts nobody: a
# path robots disallows must not enter the archive by way of an allowed URL that
# redirects to it, and the withheld address is recorded like any other.
import urllib.robotparser as _rprs
_rp = _rprs.RobotFileParser()
_rp.parse(["User-agent: *", "Disallow: /private"])
_rob = _no_net_scraper(None)
_rob._robots_by_host = {"example.com": _rp}
_rob._fetch_page = lambda url, depth: {
    "url": ("https://example.com/private" if url.endswith("/s") else url),
    "title": url, "content": [], "links": ["https://example.com/s"],
    "binary_references": [],
}
assert [q["url"] for q in _rob.scrape()] == ["https://example.com/"], _rob.visited
assert _rob.robots_excluded == ["https://example.com/private"], _rob.robots_excluded
assert _rob.total_errors == 0, _rob.total_errors


# ── A heading's own anchor is furniture ───────────────────────────────────────
# Every documentation generator appends a "link to this heading" anchor, and the
# heading is read with get_text(), so the anchor's character became part of the
# heading. Live: 225 of 225 headings on one Django settings page ended in \u00b6, and
# MkDocs-Material's Font-Awesome glyph prints in Word as an empty box.
def _page_and_headings(body_html, title="T"):
    _h = (f"<html><head><title>{title}</title></head><body><main>{body_html}"
          "<p>Body text with well over eighty characters, so the content-area "
          "selector picks this main element.</p></main></body></html>")
    _p = PageExtractor().extract(_h, "https://example.com/")
    return _p, [b["text"] for b in _p["content"] if b["type"] == "heading"]

for _anchor in ('<a class="headerlink" href="#intro" title="Link to this heading">\u00b6</a>',
                '<a class="headerlink" href="#intro">\uf0c1</a>',   # MkDocs-Material
                '<a class="hash-link" href="#intro">\u200b</a>',    # Docusaurus / VitePress
                '<a class="anchor" href="#intro">#</a>',        # GitHub-flavoured
                '<a class="anchor" href="#intro"></a>'):        # icon-only, svg already gone
    _, _hs = _page_and_headings(f"<h2>Intro{_anchor}</h2>")
    assert _hs == ["Intro"], (_anchor, _hs)

# ...and the title comes off the same h1, so it is fixed by the same pass
_pg, _ = _page_and_headings('<h1>Getting Started<a class="headerlink" href="#">\u00b6</a></h1>')
assert _pg["title"] == "Getting Started", _pg["title"]

# The guard is "no word character", never a class name: matching headerlink would
# miss hash-link and anchor, and would still take out the round-12 case below.
# A heading that is wholly a real link keeps every word of it.
_, _hs = _page_and_headings('<h3><a href="/p/1">My Post</a></h3>')
assert _hs == ["My Post"], _hs
# Mutation guard for round 12: docutils wraps the heading's own words, so its
# anchor is full of word characters and must survive. PEP 8 keeps its 42 titles.
_, _hs = _page_and_headings('<h2><a class="toc-backref" href="#x">Introduction</a></h2>')
assert _hs == ["Introduction"], _hs
# A heading of nothing but an anchor character was never a heading; emitting an
# empty one would put a blank Heading 1 in Word.
_, _hs = _page_and_headings('<h2><a class="headerlink" href="#x">\u00b6</a></h2>')
assert _hs == [], _hs

# A permalink is not a heading habit. Sphinx hangs one off every <dt> API entry
# (all 68 built-in functions on one CPython page) and every figure caption, so
# the rule is placement-blind: any same-page fragment link carrying no word.
_dt = _page_and_headings('<dl><dt>abs(x)<a class="headerlink" href="#abs">\u00b6</a></dt>'
                '<dd><p>Return the absolute value.</p></dd></dl>')[0]
assert not any("\u00b6" in (b.get("text") or "") for b in _dt["content"]), _dt["content"]
assert any("abs(x)" in (b.get("text") or "") for b in _dt["content"]), _dt["content"]

# Two guards keep the rule from eating content. A link out of the page is not a
# self-link, so a label-less arrow to the next page survives...
_arrow = _page_and_headings('<p>Read on <a href="/next">→</a></p>')[0]
assert any("\u2192" in (b.get("text") or "") for b in _arrow["content"]), _arrow["content"]
# ...and an anchor around a picture reads as empty text but is a picture. Losing
# it would drop every thumbnail that links to its own lightbox.
_lightbox = _page_and_headings('<p><a href="#modal"><img src="/i.png" alt="Diagram"></a></p>')[0]
assert [b["alt"] for b in _lightbox["content"] if b["type"] == "image"] == ["Diagram"],     _lightbox["content"]


# ── A page header holding the page's own <h1> ─────────────────────────────────
# MediaWiki's Vector 2022 puts <h1 id="firstHeading"> inside the page <header>,
# which round 12 drops as furniture everywhere else, so the title fell back to
# <title> and every page read "Python (programming language) - Wikipedia".
_wiki = ('<html><head><title>Python (programming language) - Wikipedia</title>'
         '</head><body><main><header class="vector-page-titlebar">'
         '<h1>Python (programming language)</h1></header>'
         '<p>Body text with well over eighty characters, so the content-area '
         'selector picks this main element.</p></main></body></html>')
assert PageExtractor().extract(_wiki, "https://w/")["title"] ==     "Python (programming language)"
# Round 7's guard, which this must not break: a site name in the page header is
# not the page's title, and <title> does not open with it, so <title> stays whole
# — including the "Ports 80 - 443 explained" shape a separator strip would cut.
_acme = ('<html><head><title>Install guide - ACME Docs</title></head><body><main>'
         '<header><h1>ACME Corp</h1></header>'
         '<p>Body text with well over eighty characters, so the content-area '
         'selector picks this main element.</p></main></body></html>')
assert PageExtractor().extract(_acme, "https://a/")["title"] ==     "Install guide - ACME Docs"

# The other title order, which "opens with the h1" cannot tell from the real
# thing: a site name opens "ACME Docs | Install guide" too, so every page of
# that site was titled "ACME Docs" and its whole contents list said so N times.
# No <main> here, deliberately — the content area falls back to <body>, which is
# what puts a body-level page header's <h1> in reach of early_h1.
_first = ('<html><head><title>ACME Docs | Install guide</title></head><body>'
          '<header><h1>ACME Docs</h1></header><div><h2>Requirements</h2>'
          '<p>Body text with well over eighty characters of real page content, '
          'so nothing here reads as sparse.</p></div></body></html>')
assert PageExtractor().extract(_first, "https://a/")["title"] == "ACME Docs | Install guide"

# The cost of that rule, pinned so it reads as a decision and not as a bug: a
# short page name under a long site name keeps the whole <title>. Round 13 gave
# "Go" here. Nothing on one page separates this from the ACME shape above, and
# a verbose title is per-page distinct where a repeated one is not — so if a
# later round means to change this, it changes both assertions together.
_short = ('<html><head><title>Go - Wikipedia</title></head><body><main>'
          '<header class="vector-page-titlebar"><h1>Go</h1></header>'
          '<p>Body text with well over eighty characters, so the content-area '
          'selector picks this main element.</p></main></body></html>')
assert PageExtractor().extract(_short, "https://w/")["title"] == "Go - Wikipedia"


# ── A .xml.gz sitemap is a gzip file, not a gzip Content-Encoding ─────────────
# requests decodes Content-Encoding, never a compressed body, so the XML parser
# read nothing. www.gnu.org's sitemap index points only at sitemap0.xml.gz, so
# that site's whole sitemap counted as zero pages.
import gzip as _gzip

class _SitemapResp:
    # iter_content and close: the download is read to a cap now, not just the
    # decompression, so a plain 2 GB sitemap no longer decides the run's memory
    def __init__(self, body): self.content = body

    def iter_content(self, n):
        for i in range(0, len(self.content), n):
            yield self.content[i:i + n]

    def close(self): pass

_XML = (b'<?xml version="1.0"?><urlset '
        b'xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
        b"<url><loc>https://x/a</loc></url><url><loc>https://x/b</loc></url>"
        b"</urlset>")
_an = SiteAnalyzer.__new__(SiteAnalyzer)
_an.host = "x"   # analyze() sets this; __new__ skips it, and the guard is fail-closed
for _body, _want in ((_XML, 2), (_gzip.compress(_XML), 2), (b"not xml at all", 0)):
    _an.session = type("S", (), {"get": staticmethod(
        lambda url, timeout=None, stream=None, _b=_body: _SitemapResp(_b))})()
    assert _an._parse_sitemap("https://x/sitemap.xml") == (
        ["https://x/a", "https://x/b"] if _want else []), _body[:20]

# ...and the decompression is bounded, because a sitemap is untrusted input. An
# unbounded gzip.decompress lets the site decide how much memory the run uses:
# 199 KB of gzipped zeros measured at 438 MB peak. Held to the cap here by URL
# count, which also pins that a truncated file still yields what parsed.
import crawl as _crawl_mod

_MANY = (b'<?xml version="1.0"?><urlset '
         b'xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
         + b"".join(b"<url><loc>https://x/page%05d</loc></url>" % i
                    for i in range(2000)) + b"</urlset>")
_an.session = type("S", (), {"get": staticmethod(
    lambda url, timeout=None, stream=None: _SitemapResp(_gzip.compress(_MANY)))})()
assert len(_an._parse_sitemap("https://x/sitemap.xml.gz")) == 2000   # under the cap
_was_cap = _crawl_mod.SITEMAP_MAX_BYTES
_crawl_mod.SITEMAP_MAX_BYTES = 4096
try:
    _capped = _an._parse_sitemap("https://x/sitemap.xml.gz")
finally:
    _crawl_mod.SITEMAP_MAX_BYTES = _was_cap
assert 0 < len(_capped) < 2000, len(_capped)


# ── A sitemap is only the analysed site's own to declare ──────────────────────
# Nothing checked the host of a <sitemap><loc> or of a robots.txt Sitemap: line,
# so the archived site chose what this machine fetched: 127.0.0.1:8080 and
# 169.254.169.254 both reached, measured, from --analyze-only alone.
from crawl import SITEMAP_MAX_FETCHES


class _IdxSession:
    """Serves one sitemap index, then an empty urlset for anything else."""

    def __init__(self, locs):
        self.locs = locs
        self.fetched = []

    def get(self, url, **kw):
        self.fetched.append(url)
        if url.endswith("/sitemap.xml"):
            body = (b'<?xml version="1.0"?><sitemapindex '
                    b'xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
                    + b"".join(b"<sitemap><loc>%s</loc></sitemap>" % l.encode()
                               for l in self.locs)
                    + b"</sitemapindex>")
        else:
            body = (b'<?xml version="1.0"?><urlset '
                    b'xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
                    b"<url><loc>https://x/p</loc></url></urlset>")
        return _SitemapResp(body)


_idx = _IdxSession([
    "http://127.0.0.1:8080/admin/sitemap.xml",     # loopback
    "http://169.254.169.254/latest/sitemap.xml",   # cloud metadata
    "https://evil.example/sitemap.xml",            # any third party
    "https://www.x/child-sitemap.xml",             # the site's own, www spelling
    "https://x/other-sitemap.xml",                 # the site's own
])
_an2 = SiteAnalyzer.__new__(SiteAnalyzer)
_an2.host, _an2.session = "x", _idx
_an2._sitemap_fetches = SITEMAP_MAX_FETCHES
_an2._parse_sitemap("https://x/sitemap.xml")
assert _idx.fetched == ["https://x/sitemap.xml",
                        "https://www.x/child-sitemap.xml",
                        "https://x/other-sitemap.xml"], _idx.fetched

# ── The budget bounds fetches, which is not what SITEMAP_MAX_URLS bounds ──────
# SITEMAP_MAX_URLS counts URLs collected, and a child sitemap returning none
# never advances that count: an index listing 50000 empty children performed
# 50001 fetches. This is the check that fails if the two are ever conflated again.
class _EmptyChildSession(_IdxSession):
    def get(self, url, **kw):
        self.fetched.append(url)
        if url.endswith("/sitemap.xml"):
            return _SitemapResp(
                b'<?xml version="1.0"?><sitemapindex '
                b'xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
                + b"".join(b"<sitemap><loc>https://x/s%d.xml</loc></sitemap>" % i
                           for i in range(50000))
                + b"</sitemapindex>")
        return _SitemapResp(b"<urlset/>")   # nothing, so no URL count advances


_flood = _EmptyChildSession([])
_an3 = SiteAnalyzer.__new__(SiteAnalyzer)
_an3.host, _an3.session = "x", _flood
_an3._sitemap_fetches = SITEMAP_MAX_FETCHES
assert _an3._parse_sitemap("https://x/sitemap.xml") == []
assert len(_flood.fetched) <= SITEMAP_MAX_FETCHES, len(_flood.fetched)

# ── robots.txt, the sitemap and the homepage are all bounded bodies ───────────
# read_capped already governed images and pages. These three did not, so a 2 GB
# sitemap, an endless robots.txt or a huge homepage each decided the run's memory.
import crawl as _cm


class _BodyResp:
    def __init__(self, body, url="https://x/", ct="text/html"):
        self.content, self.url = body, url
        self.status_code, self.headers = 200, {"content-type": ct}

    def iter_content(self, n):
        for i in range(0, len(self.content), n):
            yield self.content[i:i + n]

    def close(self): pass

    def raise_for_status(self): pass


_ROBOTS = b"User-agent: *\nDisallow: /private\n"
_rx2 = WebScraper.__new__(WebScraper)
_rx2._robots_by_host = {}
_rx2.session = type("S", (), {"get": staticmethod(
    lambda url, **kw: _BodyResp(_ROBOTS))})()
assert _rx2._fetch_robots("https://x") is not None, "a normal robots.txt still parses"
_was = _cm.ROBOTS_MAX_BYTES
_cm.ROBOTS_MAX_BYTES = 8
try:
    assert _rx2._fetch_robots("https://x") is None, "an oversized robots.txt reads as absent"
finally:
    _cm.ROBOTS_MAX_BYTES = _was

_an4 = SiteAnalyzer.__new__(SiteAnalyzer)
_an4.host = "x"
_an4._sitemap_fetches = SITEMAP_MAX_FETCHES
_an4.session = type("S", (), {"get": staticmethod(
    lambda url, **kw: _BodyResp(_XML, ct="application/xml"))})()
assert len(_an4._parse_sitemap("https://x/sitemap.xml")) == 2
_was = _cm.SITEMAP_MAX_DOWNLOAD_BYTES
_cm.SITEMAP_MAX_DOWNLOAD_BYTES = 8
try:
    _an4._sitemap_fetches = SITEMAP_MAX_FETCHES
    assert _an4._parse_sitemap("https://x/sitemap.xml") == [], "an oversized sitemap yields none"
finally:
    _cm.SITEMAP_MAX_DOWNLOAD_BYTES = _was

# A homepage past the cap leaves the site *reachable* — the host answered. Calling
# it unreachable sent the agent down SKILL.md's stop-here branch instead.
_HOME = b"<html><body><main><p>" + b"x" * 5000 + b"</p></main></body></html>"
_an5 = SiteAnalyzer.__new__(SiteAnalyzer)
_an5.url = "https://x/"
_an5.session = type("S", (), {"get": staticmethod(
    lambda url, **kw: _BodyResp(_HOME))})()
_was = _cm.MAX_PAGE_BYTES
_cm.MAX_PAGE_BYTES = 64
try:
    _big = _an5.analyze()
finally:
    _cm.MAX_PAGE_BYTES = _was
assert _big["reachable"] is True, "the host answered, so it is not unreachable"
assert "not analysed" in (_big["error"] or ""), _big["error"]
assert _big["js_rendered"] is False and _big["homepage_links"] == 0

# ── The host anchor follows the redirect, or every sitemap is refused ─────────
# An apex that redirects (angular.io → angular.dev) serves its robots.txt from
# the new host and names its sitemaps there.
class _MovedSession:
    def get(self, url, **kw):
        if url == "https://old.example/":
            return _BodyResp(_HOME, url="https://new.example/")
        if url == "https://new.example/robots.txt":
            return _BodyResp(b"User-agent: *\nSitemap: https://new.example/sitemap.xml\n",
                             ct="text/plain")
        if url == "https://new.example/sitemap.xml":
            return _BodyResp(_XML, ct="application/xml")
        raise AssertionError(f"unexpected fetch: {url}")


_moved = SiteAnalyzer("https://old.example/")
_moved.session = _MovedSession()
_mv = _moved.analyze()
assert _moved.host == "new.example", _moved.host
assert _mv["sitemap_page_count"] == 2, "the sitemap on the redirect target was refused"

# ── A Content-Type is the site's to write, so it is not a safe HTML value ──────
# `image/png" onload="…` passed a bare startswith("image/") check, was stored in
# pages.json, and broke straight out of the img attribute the PDF builder writes.
import base64 as _b64

from extract import MIME_RE

assert MIME_RE.fullmatch("image/svg+xml") and MIME_RE.fullmatch("image/vnd.microsoft.icon")
assert not MIME_RE.fullmatch('image/png" onload="evil')
_evil = {"type": "image", "alt": "a", "src": "https://x/y.png",
         "mime_type": 'image/png" onload="evil',
         "data": _b64.b64encode(b"\x89PNG").decode()}
assert 'onload="evil' not in _block_html(_evil), "mime escaped into an attribute"

# ── Phase 1 paces its sitemap fetches, the one place it can issue many ────────
# WebScraper and ImageHandler both paced already; the analyzer did not, so a
# sitemap index fanning out to hundreds of fetches ignored --rate-limit and the
# host's Crawl-delay alike. Never before the first: one sitemap waits for nothing.
_slept.clear()
_crawl.time = types.SimpleNamespace(sleep=_slept.append)
try:
    _paced = _IdxSession(["https://x/a.xml", "https://x/b.xml", "https://x/c.xml"])
    _ap = SiteAnalyzer("https://x/", rate_limit=2.0)
    _ap.session = _paced
    _ap._parse_sitemap("https://x/sitemap.xml")
finally:
    _crawl.time = time   # restore, so any later check sleeps for real
assert len(_paced.fetched) == 4, _paced.fetched     # the index and its three children
assert _slept == [2.0, 2.0, 2.0], _slept            # four fetches, three gaps


# ...and the pace is the one the host asks for, which is why robots.txt is read
# before any sitemap is fetched rather than after
class _DelaySession:
    def get(self, url, **kw):
        if url == "https://d.example/":
            return _BodyResp(_HOME, url="https://d.example/")
        if url == "https://d.example/robots.txt":
            return _BodyResp(b"User-agent: *\nCrawl-delay: 7\n"
                             b"Sitemap: https://d.example/one.xml\n"
                             b"Sitemap: https://d.example/two.xml\n", ct="text/plain")
        return _BodyResp(_XML, ct="application/xml")


_slept.clear()
_crawl.time = types.SimpleNamespace(sleep=_slept.append)
try:
    _da = SiteAnalyzer("https://d.example/", rate_limit=1.0)
    _da.session = _DelaySession()
    _dr = _da.analyze()
finally:
    _crawl.time = time
assert _dr["sitemap_page_count"] == 2, _dr           # both sitemaps hold the same two URLs
assert _slept == [7.0], _slept   # the host's 7s, not the caller's 1s, and only between

# ── A rebuild carries the gaps the crawl recorded ─────────────────────────────
# The References section names the domains that refused and the pages robots.txt
# withheld, so a reader months later can tell a gap in coverage from a site that
# had nothing there. A --pages-file rebuild passed neither list, so the rebuilt
# archive silently claimed complete coverage.
from main import coverage_path, read_coverage


def _gap_scraper(blocked, excluded):
    class S:
        def __init__(self, base_url, **kw):
            self.blocked_domains, self.robots_excluded = set(blocked), list(excluded)
            self.total_errors = 0

        def scrape(self, sitemap_urls=None):
            return [{"url": "https://example.com/a", "title": "A", "breadcrumb": [],
                     "depth": 0, "binary_references": [],
                     "content": [{"type": "paragraph", "text": "y" * 400}]}]
    return S


_crawl.SiteAnalyzer = _PlainAnalyzer
_crawl.WebScraper = _gap_scraper({"down.example"}, ["https://example.com/private/p"])
_gd = _tf.mkdtemp()
sys.argv = ["main.py", "--url", "https://example.com", "--format", "md",
            "--output", os.path.join(_gd, "g.md"), "--work-dir", _gd]
with contextlib.redirect_stdout(io.StringIO()):
    _main.main()
_crawl.SiteAnalyzer, _crawl.WebScraper = SiteAnalyzer, WebScraper

_pf = os.path.join(_gd, "pages.json")
assert read_coverage(_pf) == (["down.example"], ["https://example.com/private/p"])

sys.argv = ["main.py", "--url", "https://example.com", "--format", "md",
            "--output", os.path.join(_gd, "rebuilt.md"), "--pages-file", _pf]
with contextlib.redirect_stdout(io.StringIO()):
    _main.main()
for _name in ("g.md", "rebuilt.md"):
    _doc = io.open(os.path.join(_gd, _name), encoding="utf-8").read()
    assert "down.example" in _doc, _name
    assert "https://example.com/private/p" in _doc, _name

# A pages file that travelled without its sidecar still rebuilds: an absent
# record of the gaps is not a reason to refuse the pages themselves
os.remove(coverage_path(_pf))
assert read_coverage(_pf) == ([], [])
sys.argv = ["main.py", "--url", "https://example.com", "--format", "md",
            "--output", os.path.join(_gd, "bare.md"), "--pages-file", _pf]
with contextlib.redirect_stdout(io.StringIO()):
    _main.main()
assert os.path.exists(os.path.join(_gd, "bare.md"))

# -- A path prefix is judged on the normalized path ----------------------------
# _in_scope read the raw path while the visited set read the normalized one, so
# the two gates disagreed about the same page: a site links its section index as
# /docs as often as /docs/, and --include-path docs/ dropped it while
# --exclude-path docs/legacy/ archived the branch's index it was withholding.
from crawl import _path_prefixes


def _scoped(include, exclude, scope="same", allowed=()):
    w = WebScraper.__new__(WebScraper)
    w.scope, w.base_domain, w.allowed_domains = scope, "x.example", list(allowed)
    w.include_paths = _path_prefixes(include)
    w.exclude_paths = _path_prefixes(exclude)
    return w


_sec = _scoped(["docs/"], ["docs/legacy/"])
for _u, _want in [("https://x.example/docs/", True),
                  ("https://x.example/docs", True),           # the section index
                  ("https://x.example/docs/a", True),
                  ("https://x.example/docs/legacy", False),   # the excluded index
                  ("https://x.example/docs/legacy/a", False),
                  ("https://x.example/documentation", False),
                  ("https://x.example/docs/legacy.html", True)]:  # a page, not the branch
    assert _sec._in_scope(_u) is _want, (_u, _want)

# The domain is still judged on the raw netloc: _normalize folds www., and
# folding it here would leave --allowed-domains www.foo.com matching nothing
_cust = _scoped([], [], scope="custom", allowed=["www.foo.com"])
assert _cust._in_scope("https://www.foo.com/x"), "www folding leaked into the domain test"


# -- A cell spanning down outlives a row shorter than it -----------------------
# The row loop broke as soon as a <tr> ran out of its own cells, even with a
# rowspan still pending further right: that row lost the value, and the carry
# was then spent by a row the cell never covered.
_span_html = ("<table>"
              "<tr><td>a1</td><td>b1</td><td rowspan=3>R</td></tr>"
              "<tr><td>a2</td></tr>"
              "<tr><td>a3</td><td>b3</td></tr>"
              "<tr><td>a4</td><td>b4</td><td>c4</td></tr></table>")
assert PageExtractor()._table_rows(
    BeautifulSoup(_span_html, "lxml").find("table"), [], "https://x/") == [
    ["a1", "b1", "R"], ["a2", "", "R"], ["a3", "b3", "R"], ["a4", "b4", "c4"]]


# -- A page does not choose what this machine fetches --------------------------
# Round 17 gave the sitemap a host rule; an <img src> is the same untrusted
# input and had none, so an archived page could point the download at
# 127.0.0.1 or 169.254.169.254 and have the crawl probe its operator's network.
import socket as _sock

from common import host_is_internal
from extract import ImageHandler

for _h in ("127.0.0.1", "169.254.169.254", "10.0.0.1", "192.168.1.7", "::1",
           "0.0.0.0", "fe80::1%eth0", ""):
    assert host_is_internal(_h), _h
for _h in ("93.184.216.34", "8.8.8.8", "2001:4860:4860::8888"):
    assert not host_is_internal(_h), _h

# A name is resolved, not merely read - a literal-IP check alone is one DNS
# record away from useless. Any address of the name counts.
_GAI = {
    "inside.example": [(2, 1, 6, "", ("10.1.2.3", 0))],
    "mixed.example": [(2, 1, 6, "", ("93.184.216.34", 0)),
                      (2, 1, 6, "", ("127.0.0.1", 0))],
    "cdn.example": [(2, 1, 6, "", ("93.184.216.34", 0))],
}


def _fake_gai(host, port, *a, **k):
    if host not in _GAI:
        raise _sock.gaierror(host)
    return _GAI[host]


import common as _common

host_is_internal.cache_clear()
_real_gai = _common.socket.getaddrinfo
_common.socket.getaddrinfo = _fake_gai
try:
    assert host_is_internal("inside.example")
    assert host_is_internal("mixed.example"), "one private address is enough"
    assert not host_is_internal("cdn.example"), "a real CDN must keep working"
    assert not host_is_internal("nx.example"), "a name that does not resolve is not internal"
finally:
    _common.socket.getaddrinfo = _real_gai
    host_is_internal.cache_clear()


class _NoCallSession:
    def get(self, url, **kw):
        raise AssertionError("a request reached an internal address: " + url)


_guard = ImageHandler(_NoCallSession(), os.path.join(_tf.mkdtemp(), "im"))
# Refused before the request, which is the whole point of the guard
assert _guard.fetch("http://127.0.0.1:8080/logo.png", "https://public.example/") is None
assert _guard.fetch("/logo.png", "http://169.254.169.254/") is None

_PNG = b"\x89PNG\r\n\x1a\n" + b"0" * 32
_local = ImageHandler(
    type("S", (), {"get": staticmethod(
        lambda url, **kw: _BodyResp(_PNG, ct="image/png"))})(),
    os.path.join(_tf.mkdtemp(), "im"), 0.0, allow_internal=True)
# ...but a docs server archived on localhost keeps its own pictures
_got = _local.fetch("http://127.0.0.1:8000/logo.png", "http://127.0.0.1:8000/")
assert _got and _got["mime_type"] == "image/png", _got
assert io.open(_got["path"], "rb").read() == _PNG


# ...and the guard survives a redirect, which is what makes it a guard. Letting
# requests follow the hops means the request has already reached the internal
# address by the time the final URL comes back to be checked: any host a page
# links could hand the download on to 169.254.169.254 with a single 302.
class _RedirSession:
    def __init__(self, target):
        self.target, self.reached = target, []

    def get(self, url, **kw):
        self.reached.append(url)
        if url != "https://cdn.example/pic.png":
            return _BodyResp(_PNG, url=url, ct="image/png")
        hop = _BodyResp(b"", url=url)
        hop.status_code = 302
        hop.headers["location"] = self.target
        return hop


_evil = _RedirSession("http://169.254.169.254/latest/meta-data/img")
assert ImageHandler(_evil, os.path.join(_tf.mkdtemp(), "im")).fetch(
    "https://cdn.example/pic.png", "https://public.example/") is None
assert _evil.reached == ["https://cdn.example/pic.png"], \
    f"the redirect was followed to the internal address: {_evil.reached}"

# A redirect between public hosts is ordinary CDN behaviour and still resolves
_ok = _RedirSession("https://img.example/real.png")
assert ImageHandler(_ok, os.path.join(_tf.mkdtemp(), "im")).fetch(
    "https://cdn.example/pic.png", "https://public.example/") is not None
assert _ok.reached == ["https://cdn.example/pic.png",
                       "https://img.example/real.png"], _ok.reached

# The crawler sets the flag from the site it was pointed at, not per picture
assert WebScraper("http://127.0.0.1:8000/", work_dir=_tf.mkdtemp()).image_handler.allow_internal
assert not WebScraper("https://93.184.216.34/", work_dir=_tf.mkdtemp()).image_handler.allow_internal


# -- The sitemap is read only where it is used ---------------------------------
# scrape() seeds from a sitemap only on an unlimited-depth run, so a
# depth-limited build spent up to SITEMAP_MAX_FETCHES paced fetches on a list it
# then ignored - an hour of Phase 1 on a host asking 30s a request. One rule,
# stated in the analyzer rather than twice.
class _SmSession:
    robots = b"User-agent: *\nSitemap: https://g.example/sitemap.xml\n"

    def __init__(self):
        self.fetched = []

    def get(self, url, **kw):
        self.fetched.append(url)
        if url == "https://g.example/":
            return _BodyResp(_HOME, url="https://g.example/")
        if url.endswith("/robots.txt"):
            return _BodyResp(self.robots, ct="text/plain")
        return _BodyResp(_XML, ct="application/xml")


_g1 = SiteAnalyzer("https://g.example/")
_g1.session = _SmSession()
assert _g1.analyze()["sitemap_page_count"] == 2, "reading it still works"

_g0 = SiteAnalyzer("https://g.example/")
_g0.session = _SmSession()
_r0 = _g0.analyze(read_sitemap=False)
assert _r0["reachable"] and _r0["has_sitemap"], _r0   # robots.txt still declares it
assert _r0["sitemap_page_count"] == 0, _r0
assert not any("sitemap" in u for u in _g0.session.fetched), _g0.session.fetched


# ...and the existence probe is skipped too, not only the parse
class _BareSession(_SmSession):
    robots = b"User-agent: *\n"


_g2 = SiteAnalyzer("https://g.example/")
_g2.session = _BareSession()
assert not _g2.analyze(read_sitemap=False)["has_sitemap"]
assert not any("sitemap" in u for u in _g2.session.fetched), _g2.session.fetched
# The probe reads one chunk, not the whole file: it used to download up to 50MB,
# throw it away, and leave _parse_sitemap to fetch the same file again
_g3 = SiteAnalyzer("https://g.example/")
_g3.session = _BareSession()
assert _g3.analyze()["sitemap_page_count"] == 2, "the probe still finds it"


# ---------------------------------------------------------------------------
# Reading direction is a property of a block, not of the document: an archive
# carries the site's language beside an English scaffold, so one direction for
# the whole file is wrong for one half of it.
from common import is_rtl

assert is_rtl("\u0648\u0627\u062c\u0647\u0629 \u0628\u0631\u0645\u062c\u0629"), "Arabic"
assert is_rtl("\u05de\u05d3\u05e8\u05d9\u05da \u05d4\u05ea\u05e7\u05e0\u05d4"), "Hebrew"
assert not is_rtl("Installation guide"), "English"
assert not is_rtl("pip install requests"), "a command is not right-to-left"
assert not is_rtl(""), "nothing to judge"
# Predominance, not the first strong character: this heading opens in Latin and
# still belongs to an Arabic page, which the first-character rule gets wrong.
assert is_rtl("API \u2014 \u0648\u0627\u062c\u0647\u0629 \u0628\u0631\u0645\u062c\u0629 \u0627\u0644\u062a\u0637\u0628\u064a\u0642\u0627\u062a")
# ...and the inverse: one Arabic product name inside an English sentence does not
# flip the paragraph.
assert not is_rtl("The \u0648\u0627\u062c\u0647\u0629 endpoint returns a list of active users")

_AR = "\u0645\u0631\u062d\u0628\u0627 \u0628\u0627\u0644\u0639\u0627\u0644\u0645"
_rtl_blocks = [
    {"type": "paragraph", "text": _AR},
    {"type": "heading", "level": 1, "text": _AR},
    {"type": "list", "items": [_AR, _AR]},
    {"type": "table", "data": [[_AR, _AR]]},
    {"type": "callout", "text": _AR},
]
for _b in _rtl_blocks:
    assert 'dir="rtl"' in _block_html(_b), _b["type"]
# A code block keeps its left-to-right layout on every page, and English text is
# never marked: the scaffold is English, so marking it would be the same bug the
# other way round.
assert 'dir="rtl"' not in _block_html({"type": "code", "text": _AR})
assert 'dir="rtl"' not in _block_html({"type": "paragraph", "text": "Hello world"})
assert 'dir="rtl"' not in _block_html({"type": "table", "data": [["Name", "Value"]]})

# The docx side, where direction lives in w:bidi on the paragraph and w:bidiVisual
# on the table. Word shapes the glyphs without them; what they fix is the layout.
try:
    from docx import Document

    from build_docx import build_docx

    _rtl_page = {
        "title": _AR, "url": "https://x.example/a", "breadcrumb": [_AR, _AR],
        "content": [
            {"type": "heading", "level": 1, "text": _AR},
            {"type": "paragraph", "text": _AR},
            {"type": "list", "items": [_AR, _AR]},
            {"type": "table", "data": [[_AR, _AR], [_AR, _AR]]},
            {"type": "callout", "text": _AR},
            {"type": "code", "text": "pip install requests"},
        ],
    }
    _ltr_page = {
        "title": "Overview", "url": "https://x.example/b", "breadcrumb": ["Home"],
        "content": [{"type": "paragraph", "text": "Plain English paragraph here"}],
    }
    _dxr = os.path.join(_tf.mkdtemp(), "rtl.docx")
    build_docx([_rtl_page, _ltr_page], _dxr, "https://x.example/")
    _dr = Document(_dxr)
    _marked = [p.text for p in _dr.paragraphs if "<w:bidi" in p._p.xml]
    # page title, breadcrumb, heading, paragraph, two list items, callout — nothing else
    assert len(_marked) == 7, _marked
    assert all(is_rtl(t) for t in _marked), _marked
    assert "bidiVisual" in _dr.element.xml, "an Arabic table has its columns mirrored"
    # Presence is not enough: w:pPr and w:tblPr children are an ordered schema sequence, and
    # Word repairs a file that gets the order wrong. python-docx does not validate it, so an
    # "element is present" assertion passes on a file Word would refuse.
    _W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    _PPR_SEQ = ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
                "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
                "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct", "topLinePunct",
                "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid",
                "spacing", "ind", "contextualSpacing", "mirrorIndents", "suppressOverlap",
                "jc", "textDirection", "textAlignment", "textboxTightWrap", "outlineLvl",
                "divId", "cnfStyle", "rPr", "sectPr", "pPrChange"]
    _TBLPR_SEQ = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
                  "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
                  "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
                  "tblDescription", "tblPrChange"]

    def _in_order(el, seq):
        ranks = [seq.index(c.tag.split("}")[-1]) for c in el
                 if c.tag.split("}")[-1] in seq]
        return ranks == sorted(ranks), [c.tag.split("}")[-1] for c in el]

    for _p in _dr.paragraphs:
        _pPr = _p._p.find(_W + "pPr")
        if _pPr is not None:
            _ok, _kids = _in_order(_pPr, _PPR_SEQ)
            assert _ok, "w:pPr children out of schema order: %s" % _kids
    for _t in _dr.tables:
        _ok, _kids = _in_order(_t._tbl.tblPr, _TBLPR_SEQ)
        assert _ok, "w:tblPr children out of schema order: %s" % _kids
    # The cover and the Table of Contents are the skill's own English text
    assert all("<w:bidi" not in p._p.xml for p in _dr.paragraphs[:8]), "scaffold stays ltr"
    assert all("<w:bidi" not in p._p.xml for p in _dr.paragraphs
               if "English" in p.text or p.text == "Overview"), "an English page stays ltr"
except ImportError:
    print("python-docx not installed - docx reading-direction check skipped")


# ── build_document dispatch ───────────────────────────────────────────────────
# The dispatch was written out twice and the rebuild's copy passed neither coverage list, so
# a rebuilt archive said nothing about the pages robots.txt withheld. Nothing tested the
# helper that fixed it, so the same mistake at a third call site would land unnoticed.
import types as _types

import main as _m

_seen = {}


def _recorder(fmt):
    def _b(pages, output, source_url, blocked_domains=None, robots_excluded=None):
        _seen[fmt] = (pages, output, source_url, blocked_domains, robots_excluded)
    return _b


_saved_mods = {k: sys.modules.get(k) for k in ("build_docx", "build_md", "build_pdf")}
try:
    for _name, _fn in (("build_docx", "build_docx"), ("build_md", "build_md"),
                       ("build_pdf", "build_pdf")):
        _stub = _types.ModuleType(_name)
        setattr(_stub, _fn, _recorder(_name))
        sys.modules[_name] = _stub

    _pages = [{"title": "T", "url": "https://e.example/", "content": []}]
    for _fmt, _mod in (("docx", "build_docx"), ("md", "build_md"), ("pdf", "build_pdf")):
        _m.build_document(_fmt, _pages, "out." + _fmt, "https://e.example/",
                          ["blocked.example"], ["https://e.example/secret"])
        assert _mod in _seen, "build_document did not reach the %s builder" % _fmt
        _p, _o, _u, _bd, _re = _seen[_mod]
        assert _bd == ["blocked.example"], "%s lost the blocked-domain list: %r" % (_fmt, _bd)
        assert _re == ["https://e.example/secret"], "%s lost the robots list: %r" % (_fmt, _re)
        assert _p is _pages and _o == "out." + _fmt and _u == "https://e.example/"

    try:
        _m.build_document("rtf", _pages, "out.rtf", "https://e.example/", [], [])
        raise AssertionError("a fourth format must raise, not fall through to a builder")
    except ValueError:
        pass
finally:
    for _k, _v in _saved_mods.items():
        if _v is None:
            sys.modules.pop(_k, None)
        else:
            sys.modules[_k] = _v

print("build_document: dispatch and coverage lists checked")


# ── LibreOffice PDF path ──────────────────────────────────────────────────────
# Strategy 1 whenever LibreOffice is installed, and unexecuted everywhere: every other PDF
# check patches _find_libreoffice to None, CI installs weasyprint's libraries only, and the
# maintainer's machine has no LibreOffice. The binary is faked, not installed — what is left
# to break is the command line, the output filename and the failure path.
import subprocess as _sp
import tempfile as _tf

import build_pdf as _bp2

try:
    import docx as _docx_mod   # noqa: F401
    _have_docx = True
except ImportError:
    _have_docx = False
    print("python-docx not installed - LibreOffice PDF path check skipped")

if _have_docx:
    _lo_calls = []

    def _fake_soffice(cmd, **kw):
        _lo_calls.append(cmd)
        outdir = cmd[cmd.index("--outdir") + 1]
        with open(os.path.join(outdir, "document.pdf"), "wb") as fh:
            fh.write(b"%PDF-1.4 fake")
        return _sp.CompletedProcess(cmd, 0, "", "")

    _pages_lo = [{"title": "Alpha", "url": "https://e.example/a", "content":
                  [{"type": "paragraph", "text": "body"}]},
                 {"title": "Beta", "url": "https://e.example/b", "content":
                  [{"type": "paragraph", "text": "body"}]}]

    _saved_run, _saved_find = _bp2.subprocess.run, _bp2._find_libreoffice
    _tmp = _tf.mkdtemp(prefix="w2d_lo_")
    try:
        _bp2.subprocess.run = _fake_soffice
        _bp2._find_libreoffice = lambda: "/usr/bin/soffice"
        _out = os.path.join(_tmp, "archive.pdf")
        _bp2.build_pdf(_pages_lo, _out, "https://e.example/",
                       blocked_domains=["blocked.example"], robots_excluded=[])
        assert _lo_calls, "build_pdf did not take the LibreOffice branch"
        _cmd = _lo_calls[0]
        assert "--headless" in _cmd and "--convert-to" in _cmd, _cmd
        assert _cmd[_cmd.index("--convert-to") + 1] == "pdf", _cmd
        assert _cmd[0] == "/usr/bin/soffice", _cmd
        assert os.path.exists(_out), "the converted pdf did not land at output_path"

        # The intermediate docx carries a plain contents list, because headless conversion
        # cannot update a TOC field. _add_static_toc is reached from nowhere else.
        _toc_seen = {}
        import build_docx as _bdx
        _saved_bd = _bdx.build_docx          # bound before the patch, or the spy calls itself

        def _spy_docx(pages, path, url, static_toc=False, **kw):
            _toc_seen["static_toc"] = static_toc
            return _saved_bd(pages, path, url, static_toc=static_toc, **kw)

        try:
            _bdx.build_docx = _spy_docx
            _lo_calls.clear()
            _bp2.build_pdf(_pages_lo, os.path.join(_tmp, "b.pdf"), "https://e.example/")
            assert _toc_seen.get("static_toc") is True,                 "the LibreOffice path must ask for a static contents list"
        finally:
            _bdx.build_docx = _saved_bd

        # soffice logs most failures to stdout; a conversion that writes nothing must raise
        # with that output, not return a missing file.
        def _silent_soffice(cmd, **kw):
            return _sp.CompletedProcess(cmd, 1, "conversion failed: no filter", "")

        _bp2.subprocess.run = _silent_soffice
        _saved_wp = _bp2._import_weasyprint
        _bp2._import_weasyprint = lambda: "ImportError: no weasyprint here"
        try:
            _bp2.build_pdf(_pages_lo, os.path.join(_tmp, "d.pdf"), "https://e.example/")
            raise AssertionError("a conversion producing no file must raise")
        except RuntimeError as _e:
            assert "no filter" in str(_e), "the raise must carry soffice's own output: %s" % _e
        finally:
            _bp2._import_weasyprint = _saved_wp
    finally:
        _bp2.subprocess.run, _bp2._find_libreoffice = _saved_run, _saved_find

    print("build_pdf: LibreOffice command line, static TOC and failure path checked")


print("test_scrape: all checks passed")
